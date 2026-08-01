import time
t0 = time.time()
log = open(r"Docs/Buch/_buildlog.txt", "w")
def L(m):
    log.write(f"{m} {round(time.time()-t0,1)}\n"); log.flush()

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

L("imports done")

md = """# Mastering the Backtester

## Kapitel 1: Test

Dies ist ein einfacher Fliesstext ohne Bild und ohne Tabelle.

- Punkt eins
- Punkt zwei

## Kapitel 2: Test

Noch ein Absatz.
"""
lines = md.splitlines()
n = len(lines)
document = Document()
L(f"doc created, n={n}")

i = 0
step = 0
while i < n:
    step += 1
    if step % 50 == 0:
        L(f"step {step} i={i}")
    stripped = lines[i].strip()
    if stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=1)
        i += 1
        continue
    if stripped.startswith("# "):
        i += 1
        continue
    if stripped.startswith("- "):
        while i < n and lines[i].strip().startswith("- "):
            document.add_paragraph(style="List Bullet").add_run(lines[i].strip()[2:].strip())
            i += 1
        continue
    if stripped == "":
        i += 1
        continue
    buf = [stripped]
    i += 1
    while i < n and lines[i].strip() != "" and not lines[i].strip().startswith(("#", "- ")):
        buf.append(lines[i].strip())
        i += 1
    document.add_paragraph(" ".join(buf).strip())
L(f"loop done after {step} steps")
document.save(r"Docs/Buch/_t1.docx")
L("SAVED")
log.close()
