from __future__ import annotations

import math
import os
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
ASSET_DIR = OUT_DIR / "generated_assets"
DOCX_PATH = OUT_DIR / "Mastering_the_Backtester.docx"
FALLBACK_DOCX_PATH = OUT_DIR / "Mastering_the_Backtester_ueberarbeitet.docx"
VERSION = "2.0"
BUILD_DATE = date.today().strftime("%d.%m.%Y")
AUTHOR = "T. Nickel"

PAGE_WIDTH_IN = 8.27
PAGE_HEIGHT_IN = 11.69
MARGIN_IN = 0.78
CONTENT_WIDTH_IN = 6.7

NAVY = RGBColor(15, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 103, 122)
INK = RGBColor(31, 41, 55)
GOLD = RGBColor(122, 90, 0)
GREEN = RGBColor(23, 114, 69)
RED = RGBColor(155, 28, 28)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF7E6"
PALE_GREEN = "EAF7F0"


SOURCES = [
    (
        "S1",
        "MetaQuotes",
        "MetaTrader 5 Help: Strategy Testing",
        "https://www.metatrader5.com/en/terminal/help/algotrading/testing",
    ),
    (
        "S2",
        "MetaQuotes",
        "MetaTrader 5 Built-in Strategy Tester",
        "https://www.metatrader5.com/en/automated-trading/strategy-tester",
    ),
    (
        "S3",
        "QuantStart",
        "Successful Backtesting of Algorithmic Trading Strategies - Part I",
        "https://www.quantstart.com/articles/Successful-Backtesting-of-Algorithmic-Trading-Strategies-Part-I/",
    ),
    (
        "S4",
        "Knowledge-Based Systems / ScienceDirect",
        "Parameter plateau in quantitative trading strategies",
        "https://www.sciencedirect.com/science/article/abs/pii/S095070512400265X",
    ),
    (
        "S5",
        "QuantInsti",
        "Walk-Forward Optimization",
        "https://blog.quantinsti.com/walk-forward-optimization-introduction/",
    ),
    (
        "S6",
        "Investopedia",
        "Currency Pairs",
        "https://www.investopedia.com/terms/c/currencypair.asp",
    ),
    (
        "S7",
        "Investopedia",
        "Trading Multiple Time Frames in FX",
        "https://www.investopedia.com/articles/forex/08/multiple-timeframe.asp",
    ),
    (
        "S8",
        "Investopedia",
        "Basics of Currency Trading",
        "https://www.investopedia.com/financial-edge/0412/the-basics-of-currency-trading.aspx",
    ),
    (
        "S9",
        "Dukascopy Bank SA",
        "Historical Data Export",
        "https://www.dukascopy.com/swiss/english/marketwatch/historical/",
    ),
    (
        "S10",
        "Dukascopy Wiki",
        "Historical Data and History Ticks",
        "https://www.dukascopy.com/wiki/en/development/strategy-api/historical-data/history-ticks/",
    ),
    (
        "S11",
        "Dukascopy Bank SA",
        "Company and Group Overview",
        "https://www.dukascopy.com/swiss/english/about/company/",
    ),
]


DE_REPLACEMENTS = [
    ("Ã„", "Ä"),
    ("Ã–", "Ö"),
    ("Ãœ", "Ü"),
    ("Ã¤", "ä"),
    ("Ã¶", "ö"),
    ("Ã¼", "ü"),
    ("ÃŸ", "ß"),
    ("hÃ¤ndisch", "händisch"),
    ("Inhaltsuebersicht", "Inhaltsübersicht"),
    ("zusammengehoerige", "zusammengehörige"),
    ("Menuepunkte", "Menüpunkte"),
    ("Menuepunkt", "Menüpunkt"),
    ("Menue", "Menü"),
    ("Oeffnet", "Öffnet"),
    ("oeffnet", "öffnet"),
    ("Oeffnen", "Öffnen"),
    ("oeffnen", "öffnen"),
    ("ueberhaupt", "überhaupt"),
    ("Ueberoptimierung", "Überoptimierung"),
    ("ueberoptimiert", "überoptimiert"),
    ("ueberoptimieren", "überoptimieren"),
    ("ueberoptimierte", "überoptimierte"),
    ("ueber", "über"),
    ("Ueber", "Über"),
    ("hinzufuegt", "hinzufügt"),
    ("fuegt", "fügt"),
    ("Fuegt", "Fügt"),
    ("fuehrt", "führt"),
    ("Fuehrt", "Führt"),
    ("ausgefuehrt", "ausgeführt"),
    ("haette", "hätte"),
    ("waere", "wäre"),
    ("waeren", "wären"),
    ("waehrend", "während"),
    ("Waehrung", "Währung"),
    ("waehrung", "währung"),
    ("waehrungen", "währungen"),
    ("Waehrungen", "Währungen"),
    ("gewaehlte", "gewählte"),
    ("gewaehlten", "gewählten"),
    ("waehlen", "wählen"),
    ("waehlt", "wählt"),
    ("laedt", "lädt"),
    ("Laedt", "Lädt"),
    ("fuer", "für"),
    ("Fuer", "Für"),
    ("Knoepfe", "Knöpfe"),
    ("koennen", "können"),
    ("koennen", "können"),
    ("koennt", "könnt"),
    ("duerfen", "dürfen"),
    ("duerfen", "dürfen"),
    ("duerfte", "dürfte"),
    ("muss", "muss"),
    ("muessen", "müssen"),
    ("Maerkte", "Märkte"),
    ("Maerkten", "Märkten"),
    ("Marktmaerkten", "Marktmärkten"),
    ("zufaellig", "zufällig"),
    ("spaeter", "später"),
    ("spaetere", "spätere"),
    ("spaeteren", "späteren"),
    ("Laeufe", "Läufe"),
    ("Laeuft", "Läuft"),
    ("laeuft", "läuft"),
    ("laeufen", "läufen"),
    ("Pruefung", "Prüfung"),
    ("pruefung", "prüfung"),
    ("Pruef", "Prüf"),
    ("pruefen", "prüfen"),
    ("prueft", "prüft"),
    ("pruefbar", "prüfbar"),
    ("geprueft", "geprüft"),
    ("Qualitaet", "Qualität"),
    ("qualitaet", "qualität"),
    ("Plausibilitaet", "Plausibilität"),
    ("Profitabilitaet", "Profitabilität"),
    ("Volatilitaet", "Volatilität"),
    ("Fragilitaet", "Fragilität"),
    ("Diversitaet", "Diversität"),
    ("Stabilitaet", "Stabilität"),
    ("stabilitaet", "stabilität"),
    ("Sensitivitaet", "Sensitivität"),
    ("sensitivitaet", "sensitivität"),
    ("Sensitivitaets", "Sensitivitäts"),
    ("sensitivitaets", "sensitivitäts"),
    ("Verstaendnis", "Verständnis"),
    ("verstaendlich", "verständlich"),
    ("Verhaeltnis", "Verhältnis"),
    ("verhaeltnis", "verhältnis"),
    ("gefaehrlich", "gefährlich"),
    ("Gefaehrlich", "Gefährlich"),
    ("haeufig", "häufig"),
    ("Haeufig", "Häufig"),
    ("aehnlich", "ähnlich"),
    ("Aehnlich", "Ähnlich"),
    ("haengt", "hängt"),
    ("abhaengig", "abhängig"),
    ("Veraenderung", "Veränderung"),
    ("veraendern", "verändern"),
    ("veraendert", "verändert"),
    ("wuerde", "würde"),
    ("zerstoeren", "zerstören"),
    ("Oberflaeche", "Oberfläche"),
    ("Oberflaechen", "Oberflächen"),
    ("Flaeche", "Fläche"),
    ("glueck", "glück"),
    ("Glueck", "Glück"),
    ("Rueckgang", "Rückgang"),
    ("Rueckgaengen", "Rückgängen"),
    ("Rueckschlaege", "Rückschläge"),
    ("Bruecke", "Brücke"),
    ("Kontogroesse", "Kontogröße"),
    ("Stichprobengroesse", "Stichprobengröße"),
    ("Zielgroesse", "Zielgröße"),
    ("Groessere", "Größere"),
    ("gross", "groß"),
    ("groessere", "größere"),
    ("hoehere", "höhere"),
    ("Hoehere", "Höhere"),
    ("hoeher", "höher"),
    ("hoechste", "höchste"),
    ("schoenste", "schönste"),
    ("schoen", "schön"),
    ("zuverlaessig", "zuverlässig"),
    ("unzuverlaessig", "unzuverlässig"),
    ("gueltig", "gültig"),
    ("vollstaendig", "vollständig"),
    ("zusaetzlich", "zusätzlich"),
    ("Zusaetzlich", "Zusätzlich"),
    ("schuetzt", "schützt"),
    ("Schuetzt", "Schützt"),
    ("verkuerzen", "verkürzen"),
    ("Suchraeume", "Suchräume"),
    ("Suchraum", "Suchraum"),
    ("Paesse", "Pässe"),
    ("Saeule", "Säule"),
    ("zaehlen", "zählen"),
    ("fuenf", "fünf"),
    ("frueh", "früh"),
    ("gehoeren", "gehören"),
    ("klaeren", "klären"),
    ("erklaert", "erklärt"),
    ("Erklaerung", "Erklärung"),
    ("erklaerung", "erklärung"),
    ("moeglich", "möglich"),
    ("auffaellig", "auffällig"),
    ("spektakulaer", "spektakulär"),
    ("staendig", "ständig"),
    ("staerkere", "stärkere"),
    ("unberuehrten", "unberührten"),
    ("unberuehrtes", "unberührtes"),
    ("schuetzen", "schützen"),
    ("anfaellig", "anfällig"),
    ("Massenlaeufe", "Massenläufe"),
    ("Vergleichslaeufe", "Vergleichsläufe"),
    ("Testlaeufe", "Testläufe"),
    ("Vorpruefung", "Vorprüfung"),
    ("Vorpruefungen", "Vorprüfungen"),
    ("Serienausfuehrung", "Serienausführung"),
    ("Vollstaendige", "Vollständige"),
    ("Spaeteres", "Späteres"),
    ("Laesst", "Lässt"),
    ("spaet", "spät"),
    ("Aendern", "Ändern"),
    ("Aenderungen", "Änderungen"),
    ("aendern", "ändern"),
    ("geaendert", "geändert"),
    ("Loeschen", "Löschen"),
    ("loeschen", "löschen"),
    ("laesst", "lässt"),
    ("haengt", "hängt"),
    ("schlaegt", "schlägt"),
    ("faellt", "fällt"),
    ("verfuegbar", "verfügbar"),
    ("Verfuegbar", "Verfügbar"),
    ("Bauchgefuehl", "Bauchgefühl"),
    ("Datenbankgestuetzte", "Datenbankgestützte"),
    ("Einschaetzung", "Einschätzung"),
    ("Ergaenzung", "Ergänzung"),
    ("Frueher", "Früher"),
    ("Gedaechtnis", "Gedächtnis"),
    ("Parameterveraenderungen", "Parameterveränderungen"),
    ("Trendlaenge", "Trendlänge"),
    ("Wuerde", "Würde"),
    ("abhaengen", "abhängen"),
    ("glaubwuerdiger", "glaubwürdiger"),
    ("haengen", "hängen"),
    ("laengeren", "längeren"),
    ("naechste", "nächste"),
    ("staerker", "stärker"),
    ("temporaer", "temporär"),
    ("verdaechtig", "verdächtig"),
    ("verhaelt", "verhält"),
    ("zusammenhaengen", "zusammenhängen"),
    ("Luecken", "Lücken"),
    ("luecken", "lücken"),
    ("nuetzlich", "nützlich"),
    ("Hintergruende", "Hintergründe"),
    ("Zeitraeume", "Zeiträume"),
    ("Raeume", "Räume"),
    ("grosse", "große"),
    ("grossen", "großen"),
    ("grosser", "großer"),
    ("Ausfuehrung", "Ausführung"),
    ("Ausfuehrungen", "Ausführungen"),
    ("ausfuehren", "ausführen"),
    ("ausfuehrt", "ausführt"),
    ("vollstaendiges", "vollständiges"),
    ("laenglich", "länglich"),
]


def de(text: str) -> str:
    if not isinstance(text, str):
        return text
    for old, new in DE_REPLACEMENTS:
        text = text.replace(old, new)
    return text


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    bold_candidates = [
        "C:/Windows/Fonts/segoeuib.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
    ]
    for path in (bold_candidates if bold else candidates):
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        current = ""
        for word in words:
            test = f"{current} {word}".strip()
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, max_width: int, fnt, fill="#0f172a", spacing=5):
    x, y = xy
    for line in wrap(draw, de(text), max_width, fnt):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + spacing


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline="#233044", radius=16, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, fill="#334155", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    pts = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)] if x2 >= x1 else [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(pts, fill=fill)


def make_manual_graphics() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    title = font(30, True)
    h = font(17, True)
    b = font(13)
    paths: dict[str, Path] = {}

    def save(name: str, img: PILImage.Image):
        path = ASSET_DIR / f"{name}.png"
        img.save(path)
        paths[name] = path

    img = PILImage.new("RGB", (1400, 760), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((56, 34), de("Backtester als Aufsatz auf MetaTrader"), font=title, fill="#0f2545")
    boxes = [
        ((80, 180, 310, 330), "#E8EEF5", "Nutzer", "waehlt EA, Symbol, Zeitraum, Parameter"),
        ((390, 180, 680, 330), "#EAF7F0", "Backtester", "JavaFX-UI, Presets, tester.ini, Prozessschutz"),
        ((760, 180, 1060, 330), "#FFF7E6", "MetaTrader", "Strategy Tester, Backtest, Optimierung, Forward"),
        ((1110, 180, 1330, 330), "#F4F6F9", "Reports", "HTML/XML, Logs, Kennzahlen, Equity"),
        ((390, 480, 680, 625), "#E0F2FE", "Auswertung", "Parser, Filter, Score, Sensitivitaet, KI"),
        ((760, 480, 1060, 625), "#F3E8FF", "Persistenz", "SQLite: Historie, Workflow, Reviews"),
    ]
    for box, fill, head, body in boxes:
        rounded(d, box, fill)
        d.text((box[0] + 18, box[1] + 20), de(head), font=h, fill="#0f172a")
        draw_wrapped(d, (box[0] + 18, box[1] + 60), body, box[2] - box[0] - 36, b)
    for s, e in [((310, 255), (390, 255)), ((680, 255), (760, 255)), ((1060, 255), (1110, 255)), ((535, 330), (535, 480)), ((910, 330), (910, 480)), ((680, 552), (760, 552))]:
        arrow(d, s, e)
    d.text((90, 675), de("Kernaussage: Das Projekt ersetzt den MetaTrader nicht. Es automatisiert ihn, wertet Reports aus und fuegt Robustheitslogik hinzu."), font=h, fill="#8a4b00")
    save("metatrader_wrapper", img)

    img = PILImage.new("RGB", (1400, 780), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((56, 34), de("Plateau-Optimierung: stabile Flaeche statt scharfer Spitze"), font=title, fill="#0f2545")
    left_origin = (110, 600)
    right_origin = (770, 600)
    for origin, label, color in [(left_origin, "Spitzen-Optimum", "#b91c1c"), (right_origin, "Stabiles Plateau", "#047857")]:
        ox, oy = origin
        d.line((ox, oy, ox + 480, oy), fill="#334155", width=3)
        d.line((ox, oy, ox, oy - 360), fill="#334155", width=3)
        pts = []
        for i in range(0, 480):
            x = i / 480
            if label.startswith("Spitzen"):
                y = 35 + 310 * math.exp(-((x - 0.52) ** 2) / 0.0025) + 35 * math.sin(x * 34)
            else:
                base = 230 if 0.28 < x < 0.74 else 70
                edge = 1 / (1 + math.exp(-45 * (x - 0.28))) * (1 / (1 + math.exp(45 * (x - 0.74))))
                y = 60 + 260 * edge + 12 * math.sin(x * 28)
            pts.append((ox + i, oy - int(y)))
        d.line(pts, fill=color, width=5)
        d.text((ox + 90, oy + 28), de(label), font=h, fill=color)
        d.text((ox + 4, oy + 8), de("Parameter"), font=b, fill="#334155")
        d.text((ox - 70, oy - 345), de("Score"), font=b, fill="#334155")
    draw_wrapped(d, (105, 120), "Gefaehrlich: Ein einzelner hoher Punkt kann historisches Rauschen sein. Kleine Parameter-Aenderungen zerstoeren das Ergebnis.", 500, h, "#7f1d1d")
    draw_wrapped(d, (770, 120), "Gesucht: Eine breite Zone mit aehnlich guten Resultaten. Hier ist die Strategie weniger abhaengig von exakt einem Parameterwert.", 520, h, "#14532d")
    d.text((115, 705), de("Im Backtester zeigt sich dieses Denken in Sensitivitaets-Sweeps, CV-Werten, Konsistenz, Forward-Filter und Step-7-Validierung."), font=h, fill="#0f2545")
    save("plateau_optimization", img)

    img = PILImage.new("RGB", (1400, 760), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((56, 34), de("Markt- und Timeframe-Screening"), font=title, fill="#0f2545")
    cols = ["M1/M5", "M15/M30", "H1/H4", "D1+"]
    rows = [
        ("Majors", "EURUSD, GBPUSD, USDJPY: liquide, oft enge Spreads"),
        ("Crosses", "EURGBP, EURJPY, GBPJPY: andere Dynamik ohne direkten USD-Fokus"),
        ("Metalle/Oel", "XAUUSD, XAGUSD, XTIUSD: andere Volatilitaet und Session-Effekte"),
    ]
    x0, y0 = 90, 165
    cw, rh = 260, 128
    for c, label in enumerate(["Marktgruppe"] + cols):
        rounded(d, (x0 + c * cw, y0, x0 + (c + 1) * cw - 10, y0 + 65), "#E8EEF5")
        d.text((x0 + c * cw + 22, y0 + 22), de(label), font=h, fill="#0f172a")
    for r, (group, desc) in enumerate(rows):
        y = y0 + 78 + r * rh
        rounded(d, (x0, y, x0 + cw - 10, y + rh - 12), "#F4F6F9")
        d.text((x0 + 20, y + 18), de(group), font=h, fill="#0f172a")
        draw_wrapped(d, (x0 + 20, y + 52), desc, cw - 50, b)
        for c in range(1, 5):
            fill = ["#FEE2E2", "#FEF3C7", "#EAF7F0", "#E0F2FE"][c - 1]
            rounded(d, (x0 + c * cw, y, x0 + (c + 1) * cw - 10, y + rh - 12), fill)
            msg = ["sehr laut", "kurzfristig", "ausgewogen", "trendnah"][c - 1]
            d.text((x0 + c * cw + 34, y + 45), de(msg), font=h, fill="#0f172a")
    d.text((95, 650), de("Der Multi-Backtester sucht keine Magie. Er prueft systematisch, ob eine Idee ueber mehrere Maerkte und Zeithorizonte stabil wirkt."), font=h, fill="#0f2545")
    save("market_timeframes", img)

    return paths


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Calibri"
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = MUTED
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("Mastering the Backtester | Professionelles Benutzerhandbuch")
    set_font(r, size=8.5, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Seite ")
    set_font(r, size=8.5, color=MUTED)
    add_page_number(footer)


def p(doc: Document, text: str, style: str | None = None, bold_start: str | None = None):
    text = de(text)
    if bold_start:
        bold_start = de(bold_start)
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r1 = para.add_run(bold_start)
        set_font(r1, bold=True, color=INK)
        r2 = para.add_run(text[len(bold_start):])
        set_font(r2, color=INK)
    else:
        r = para.add_run(text)
        set_font(r, color=INK)
    return para


def h1(doc: Document, text: str):
    return doc.add_heading(de(text), level=1)


def h2(doc: Document, text: str):
    return doc.add_heading(de(text), level=2)


def h3(doc: Document, text: str):
    return doc.add_heading(de(text), level=3)


def bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(de(text))
    set_font(r, color=INK)
    return para


def number(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(de(text))
    set_font(r, color=INK)
    return para


def numbered(doc: Document, items: list[str]):
    for idx, text in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.22)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        para.paragraph_format.space_after = Pt(5)
        r1 = para.add_run(f"{idx}. ")
        set_font(r1, bold=True, color=INK)
        r2 = para.add_run(de(text))
        set_font(r2, color=INK)


def callout(doc: Document, title: str, body: str, fill=LIGHT_GRAY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_IN])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    pp = cell.paragraphs[0]
    rr = pp.add_run(de(title))
    set_font(rr, size=10.5, bold=True, color=accent)
    pp.paragraph_format.space_after = Pt(3)
    pp2 = cell.add_paragraph()
    rr2 = pp2.add_run(de(body))
    set_font(rr2, size=10, color=INK)
    pp2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def image(doc: Document, path: Path, caption: str, width=6.35):
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Inches(width))
        if caption:
            cap = doc.add_paragraph(de(caption), style="Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p(doc, f"[Bild fehlt: {path}]")


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], title: str | None = None):
    if title:
        h3(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(de(head))
        set_font(r, size=9, bold=True, color=NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(de(val))
            set_font(r, size=8.5, color=INK)
    doc.add_paragraph()


def source_note(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    r = para.add_run(de(text))
    set_font(r, size=9, italic=True, color=MUTED)


def add_cover(doc: Document, graphics: dict[str, Path]):
    cover = ROOT / "images" / "backtester_platform.png"
    if cover.exists():
        image(doc, cover, "", width=6.5)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("Mastering the Backtester")
    set_font(r, size=30, bold=True, color=NAVY)
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(4)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(de("Professionelles Handbuch fuer Backtesting, Multi-Backtesting und robuste Strategieoptimierung"))
    set_font(r, size=13.5, color=MUTED)
    para.paragraph_format.space_after = Pt(18)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(de(f"Autor: {AUTHOR} | Version {VERSION} | Stand: {BUILD_DATE}"))
    set_font(r, size=10.5, italic=True, color=GOLD)

    callout(
        doc,
        "Leitgedanke",
        "Dieses Dokument erklaert nicht nur, welche Knoepfe es gibt. Es erklaert, warum jeder Bereich existiert: damit Optimierungen nicht auf einzelne historische Spitzen zielen, sondern robuste Parameterzonen, saubere Forward-Tests und echte Out-of-Sample-Validierung nutzen.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    doc.add_page_break()


def add_toc(doc: Document):
    h1(doc, "Inhaltsverzeichnis")
    p(doc, "Dieses Inhaltsverzeichnis ist bewusst als Arbeitskarte aufgebaut. Es zeigt nicht nur Kapitelnummern, sondern auch, welche Frage der jeweilige Abschnitt beantwortet. So kann der Leser das Dokument wie ein Lehrbuch lesen oder gezielt als Bedienhandbuch nachschlagen.")
    entries = [
        ("1", "Vorwort und Orientierung", "Zweck, Zielgruppe, MetaTrader-Rolle und Robustheitsziel."),
        ("2", "Grundlagen des Backtestings", "Historische Tests, Grenzen, Biases und Overfitting."),
        ("3", "Programmaufbau", "Arbeitsbereiche, Menuepunkte und fachlicher Ablauf."),
        ("4", "Backtester: Einzeltests", "Konfiguration, Ergebniszeile, Einzelreport und Beispielstrategie."),
        ("5", "Multi-Backtester", "Maerkte, Timeframes, Ergebniscluster und Uebertragbarkeit."),
        ("6", "Optimizer und Plateau-Auswahl", "Optimierung, Filter, Score, Evaluator und Sensitivitaet."),
        ("7", "Workflow, KI und Controlling", "Acht Stufen, Automatisierung, LLM-Analyse und OOS-Gate."),
        ("8", "Daten und Nebenbereiche", "Settings, Dukascopy, Database, Log und Manual."),
        ("9", "Parameter und Best Practices", "Nachschlagewerk fuer Einstellungen und robuste Praxis."),
        ("10", "Quellenverzeichnis", "Fachquellen, Projektbezug und Bildnachweis."),
    ]
    table(doc, ["Teil", "Kapitel", "Orientierung"], entries, [0.55, 2.55, 3.4])
    callout(
        doc,
        "Leselogik",
        "Kapitel 1 bis 3 bauen das fachliche Fundament. Kapitel 4 bis 8 sind die eigentliche Bedienungsanleitung. Kapitel 9 und 10 dienen als Nachschlagewerk, wenn man waehrend der Arbeit schnell eine Einstellung, einen Filter oder eine Quelle pruefen will.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    doc.add_page_break()


def add_frontmatter(doc: Document):
    h1(doc, "1. Vorwort und Orientierung")
    p(doc, "Dieses Handbuch ist die Bruecke zwischen zwei Welten: der praktischen Bedienung des Backtester-Programms und dem fachlichen Denken hinter robuster Strategieentwicklung. Der Leser soll nicht nur wissen, wo ein Feld liegt, sondern auch verstehen, warum dieses Feld wichtig ist und welche Fehlentscheidung es verhindern kann.")
    p(doc, "Das Projekt Backtester ist bewusst kein Ersatz fuer MetaTrader. MetaTrader bleibt die Simulationsmaschine: Er laedt historische Daten, startet Expert Advisors, erzeugt Backtest- und Optimierungsreports und fuehrt den Strategy Tester aus. Das Backtester-Projekt liegt als Aufsatz darueber. Es erzeugt Konfigurationen, startet MT4/MT5 reproduzierbar, sammelt Reports und Logfiles, liest Kennzahlen automatisch aus, speichert Ergebnisse in SQLite und fuegt Bewertungslogik hinzu, die in MetaTrader so nicht komfortabel vorhanden ist.")
    p(doc, "Das globale Ziel lautet: Strategien sollen optimiert werden, ohne in Ueberoptimierung zu kippen. Eine perfekte Equity-Kurve in der Vergangenheit ist wertlos, wenn sie nur aus einem zufaelligen Parameter-Treffer entstanden ist. Darum fuehrt das Tool vom Einzeltest ueber Multi-Backtests, Combined Analysis, Filter, Score-Gewichtung, Sensitivitaetsanalyse, KI-Bewertung und Step-7-Out-of-Sample-Validierung zu einer nachvollziehbaren Entscheidungskette.")
    callout(
        doc,
        "Was Sie nach dem Lesen koennen sollen",
        "Sie sollen erkennen, welcher Programmbereich fuer welche Frage gedacht ist: Laeuft der EA technisch? Auf welchen Maerkten wirkt die Idee stabil? Welche Parameterzone ist robust? Wie vermeidet man Optimierung auf Spitzen? Welche Kandidaten duerfen in den Best-Ordner und welche nicht?",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    h2(doc, "Kurzfassung fuer den schnellen Einstieg")
    bullet(doc, "Backtester: Einzelne MT4/MT5-Laeufe mit festen Parametern starten, Reports automatisch auswerten und technische Fehler sichtbar machen.")
    bullet(doc, "Multi-Backtester: Einen EA ueber viele Symbole und Timeframes laufen lassen, um Markt- und Zeithorizont-Eignung zu erkennen.")
    bullet(doc, "Optimizer: MetaTrader-Optimierung nutzen, aber Ergebnisse mit eigenen Filtern, Score-Gewichtung, Forward-Logik und Plateau-Denken bewerten.")
    bullet(doc, "Workflow Automator: Aus einzelnen Werkzeugen eine Anti-Curvefitting-Pipeline machen.")
    bullet(doc, "Controlling: Exportierte Strategien spaeter nachtesten, kommentieren und nicht aus dem historischen Kontext verlieren.")
    source_note(doc, "Grundlage: MetaQuotes beschreibt den Strategy Tester als Werkzeug zum Testen und Optimieren von Expert Advisors, waehrend Forward-Tests helfen sollen, Parameter-Fitting zu vermeiden [S1, S2].")
    doc.add_page_break()


def add_backtesting_foundation(doc: Document, graphics: dict[str, Path]):
    h1(doc, "2. Grundlagen: Backtesting mit MetaTrader")
    image(doc, graphics["metatrader_wrapper"], "Das Backtester-Projekt als Orchestrierungs- und Auswertungsschicht ueber MetaTrader.", width=6.45)

    h2(doc, "Backtesting in einfachen Worten")
    p(doc, "Ein Backtest ist eine historische Simulation. Eine Handelsregel wird mit Daten aus der Vergangenheit ausgefuehrt, als ob der Expert Advisor damals bereits live gehandelt haette. Das Ergebnis zeigt, welche Trades entstanden waeren, welcher Gewinn oder Verlust angefallen waere und wie gross zwischenzeitliche Rueckschlaege gewesen waeren.")
    p(doc, "Der Nutzen liegt nicht darin, die Zukunft exakt vorherzusagen. Der Nutzen liegt darin, schlechte Ideen frueh auszusortieren, technische Fehler zu finden und gute Ideen unter kontrollierten Annahmen vergleichbar zu machen. QuantStart beschreibt Backtesting als Filter-, Modellierungs-, Optimierungs- und Verifikationswerkzeug, warnt aber zugleich vor typischen Biases wie Optimierungsbias und Look-ahead Bias [S3].")
    p(doc, "MetaTrader fuehrt den eigentlichen Test aus. Das Backtester-Projekt nutzt also nicht einen selbstgebauten zweiten Markt-Simulator, sondern die offiziellen MT4/MT5-Mechanismen. Wichtig ist der Unterschied: Die Eingabeparameter stammen aus der MetaTrader-Welt, aber die Arbeitsweise danach ist erweitert. Das Tool erzeugt tester.ini-Dateien, startet Prozesse, ueberwacht Laufzeit und Logs, findet Reports, liest Kennzahlen und macht Ergebnisse vergleichbar.")

    h2(doc, "Warum automatische Auswertung wichtig ist")
    p(doc, "Ein einzelner MT5-Report kann manuell gelesen werden. Bei zehn Symbolen, fuenf Timeframes und vielen Optimierungs-Paessen wird manuelle Auswertung aber unzuverlaessig. Menschen achten schnell nur auf Profit, uebersehen geringe Tradezahlen, ignorieren Forward-Einbruch oder vergleichen Ergebnisse mit unterschiedlichen Kontoannahmen. Genau hier beginnt der Nutzen des Backtester-Projekts.")
    callout(
        doc,
        "Automatisierung bedeutet nicht Blindflug",
        "Das Programm nimmt dem Nutzer Fleissarbeit ab: starten, warten, auslesen, speichern, filtern. Die fachliche Entscheidung bleibt aber beim Nutzer. Ein Score ist ein Hinweis, kein Live-Freibrief.",
        fill=LIGHT_GRAY,
        accent=BLUE,
    )

    h2(doc, "Die zentrale Gefahr: Ueberoptimierung")
    p(doc, "Optimierung ist notwendig, weil viele Expert Advisors Eingabeparameter besitzen: Perioden, Schwellen, Stop-Loss, Take-Profit, Filter, Money-Management und weitere Regeln. Optimierung bedeutet, systematisch verschiedene Kombinationen dieser Werte zu testen. Das Problem: Wenn man genug Kombinationen probiert, findet man fast immer eine Kombination, die historisch hervorragend aussieht.")
    p(doc, "Diese Kombination kann echtes Signal enthalten. Sie kann aber auch nur historisches Rauschen getroffen haben. MetaTrader selbst bietet Optimierung, genetische Algorithmen, Forward-Tests sowie 2D/3D-Visualisierung an [S2]. Das Backtester-Projekt baut darauf auf und fragt zusaetzlich: Ist der Kandidat stabil? Bleibt er im Forward plausibel? Gibt es genug Trades? Wie reagiert der Profit, wenn Parameter leicht verschoben werden?")
    source_note(doc, "MetaQuotes: Bei der Optimierung wird eine Strategie mehrfach mit unterschiedlichen Eingabeparametern getestet; bei sehr vielen Kombinationen kann ein genetischer Algorithmus die Suche verkuerzen [S2].")


def add_program_overview(doc: Document):
    h1(doc, "3. Programmaufbau und Arbeitsbereiche")
    p(doc, "Die Menuepunkte der Anwendung sind keine zufaellige Sammlung von Tabs. Sie bilden einen Entwicklungsprozess ab. Man kann ihn in fuenf zusammengehoerige Bereiche gliedern: Vorbereitung, Einzeltest, Batch-Vergleich, Optimierung/Robustheit und Ergebnissicherung.")
    groups = [
        ["Vorbereitung", "Settings, Dukascopy Data, Log", "Pfade, Daten, Brokerzeit, MT4/MT5-Start und Fehlersuche klaeren.", "Ohne saubere Daten und Pfade sind alle folgenden Ergebnisse zweifelhaft."],
        ["Einzeltest", "Backtest", "Einen EA mit festem Parameter-Set auf Symbol und Zeitraum pruefen.", "Technische Plausibilitaet vor Batch oder Optimierung."],
        ["Batch-Vergleich", "Multi-Backtester", "Viele Symbol/Timeframe-Kombinationen mit denselben Annahmen testen.", "Erkennt, auf welchen Maerkten eine Idee ueberhaupt Sinn macht."],
        ["Optimierung", "Optimizer, Robustness, Workflow Automator", "Suchraeume, Forward, Filter, Sensitivitaet, KI und OOS-Validierung.", "Schuetzt vor Ueberoptimierung auf historische Spitzen."],
        ["Ergebnissicherung", "Database, Controlling, Manual", "Historie, Reviews, Nachtests, Export und Dokumentation.", "Strategien bleiben nachvollziehbar und pruefbar."],
    ]
    table(doc, ["Bereich", "Menuepunkte", "Aufgabe", "Warum wichtig"], groups, [1.25, 1.65, 2.0, 1.6], "Die Programmbereiche im Zusammenhang")

    h2(doc, "Die Menuepunkte sauber erklaert")
    p(doc, "Backtest ist die Werkbank fuer einzelne Laeufe. Multi-Backtester ist die Markt- und Timeframe-Matrix. Optimizer ist die Analysezentrale fuer Parameter. Robustness ist die freie Stresstest-Werkbank. Workflow Automator fuehrt die robuste Pipeline als Zustand. Controlling ist der Ort fuer spaetere Bewertung und Nachtests. Database bewahrt die Historie. Dukascopy Data verbessert die Datenversorgung. Settings setzt die technischen Grundlagen. Log zeigt, was im Hintergrund wirklich passiert. Manual ist die eingebaute Schnellhilfe.")
    menu_rows = [
        ["Backtest", "Einzelner MT4/MT5-Lauf mit festen Parametern.", "Wenn ein SET technisch geprueft, ein Fehler gesucht oder ein finaler Kandidat nachgetestet werden soll."],
        ["Multi-Backtester", "Serielle Batch-Laeufe ueber Symbole und Perioden.", "Wenn eine Idee zuerst ueber Maerkte und Timeframes gesichtet werden soll."],
        ["Optimizer", "MT5-Optimierung plus Combined Analysis, Filter, Score, Sensitivitaet und KI.", "Wenn ein Suchraum systematisch und robust bewertet werden soll."],
        ["Robustness", "Gezielte Robustheits- und Sensitivitaetspruefung einzelner Kandidaten.", "Wenn man eine ausgewaehlte Strategie gegen Parameter- und Zeitverschiebungen testen will."],
        ["Workflow Automator", "Gefuehrte Anti-Curvefitting-Pipeline mit Zustand und Gates.", "Wenn ein Kandidat professionell vom Setup bis zur OOS-Validierung gefuehrt werden soll."],
        ["Controlling", "Nachtest, Review und Pflege exportierter Strategien.", "Wenn Strategien nach dem Export nicht vergessen, sondern kontrolliert werden sollen."],
        ["Database / History", "Persistierte Backtests, Optimierungen, KI-Berichte und Zustandsdaten.", "Wenn Ergebnisse spaeter nachvollziehbar bleiben muessen."],
        ["Dukascopy Data", "Download, Scan, Konvertierung und MT5-Import externer Tickdaten.", "Wenn die Broker-Historie nicht als alleinige Datenbasis reichen soll."],
        ["Settings", "Terminalpfade, Portable Mode, Ordner, Konto-Defaults, Zeitzone.", "Wenn die technische Grundlage fuer reproduzierbare Laeufe gesetzt wird."],
        ["Log / Manual", "Laufzeitdiagnose und eingebaute Hilfe.", "Wenn etwas nicht laeuft oder eine Funktion schnell nachgeschlagen werden soll."],
    ]
    table(doc, ["Menuepunkt", "Was dort passiert", "Wann man ihn benutzt"], menu_rows, [1.35, 2.65, 2.5], "Vollstaendige Menue-Orientierung")
    callout(
        doc,
        "Leselogik fuer Anwender",
        "Wer neu beginnt, sollte nicht mit dem Optimizer starten. Erst Settings und Daten pruefen, dann einen Einzeltest laufen lassen, danach mit dem Multi-Backtester Maerkte sondieren und erst dann optimieren.",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def add_backtester_section(doc: Document):
    h1(doc, "4. Backtester: Einzeltests verstehen und auswerten")
    image(doc, ROOT / "images" / "backtester_ui1.png", "Backtest-Tab: oben Konfiguration, unten Historie und Ergebnisaktionen.", width=6.35)
    h2(doc, "Die Grafik lesen: oberer und unterer Arbeitsbereich")
    p(doc, "Der Screenshot zeigt den Backtest-Tab in zwei klar getrennten Bereichen. Oben liegt die Backtest Configuration. Dort wird festgelegt, welcher Expert Advisor getestet wird, auf welchem Symbol er laeuft, welcher Timeframe genutzt wird, welches Datumsfenster gilt und mit welchen Kontoannahmen MetaTrader den Test rechnen soll. Dieser obere Bereich ist also die Eingabe- und Startzone.")
    p(doc, "Im sichtbaren Beispiel ist der EA Market\\Scalper Deriv gewaehlt, das Symbol ist USDJPY, der Timeframe H1, der Zeitraum laeuft vom 02.11.2025 bis 02.05.2026, das Startkapital betraegt 10.000 USD, der Hebel steht auf 1:100 und das Tickmodell ist 1 minute OHLC. Das sind klassische MetaTrader-Testerangaben. Der Backtester baut daraus keinen eigenen Markt, sondern uebergibt diese Angaben kontrolliert an MetaTrader.")
    p(doc, "Unten liegt Backtest History & Results. Dieser Bereich ist das Ergebnisprotokoll. Jede Zeile steht fuer einen Lauf und zeigt die wichtigsten Kennzahlen: Expert, Symbol, Period, Profit, Trades, Win Rate und Drawdown. Die Schaltflaechen Open HTML Report und Open Directory fuehren danach in die Detailanalyse: entweder direkt zum Report oder in den Ordner mit den erzeugten Dateien.")
    callout(
        doc,
        "So liest man die Beispielstrategie",
        "Die gezeigte Zeile wirkt technisch erfolgreich: Profit 13,22, zwei Trades, 100,0 Prozent Win Rate und 2,11 Prozent Drawdown. Fachlich ist das aber noch keine starke Strategieaussage. Zwei Trades sind viel zu wenig, um Robustheit zu beurteilen. Die Zeile sagt vor allem: Der EA laeuft, handelt und schreibt auswertbare Ergebnisse. Fuer eine echte Freigabe braucht man laengere Zeitraeume, mehr Trades, Multi-Backtests, Sensitivitaet, Forward- und OOS-Pruefung.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    h2(doc, "Warum dieser Bereich gebraucht wird")
    p(doc, "Der Backtest-Tab beantwortet die erste praktische Frage: Laeuft dieser Expert Advisor mit diesem Parameter-Set technisch und fachlich plausibel? Bevor man hunderte Optimierungs-Paesse startet, muss ein einzelner Lauf sauber durchgehen. Der Tab ist deshalb Diagnosewerkzeug, Basispruefung und spaeter auch Nachtest-Werkbank.")
    p(doc, "Hier wird nichts Neues simuliert. Die Felder entsprechen MetaTrader-Strategy-Tester-Parametern: Expert Advisor, Symbol, Zeitraum, Kontogroesse, Waehrung, Hebel und Tickmodell. Der Mehrwert des Programms liegt in der Automatisierung: Es schreibt die Konfiguration, startet MT4/MT5, sammelt den Report und speichert die wichtigsten Kennzahlen.")
    h2(doc, "Was jedes Feld bedeutet")
    rows = [
        ["Expert Advisor", "Der zu testende EA.", "Ohne gueltigen EA kein Test. Zuerst pruefen, ob der EA in MT4/MT5 kompiliert ist."],
        ["Symbol", "Markt, z.B. EURUSD, USDJPY oder XAUUSD.", "Symbol muss in MetaTrader vorhanden sein; bei Custom Symbols vorher importieren."],
        ["Period", "Timeframe des Charts, z.B. M15, H1, D1.", "Timeframe muss zur Strategieidee passen."],
        ["Dates", "Historischer Testzeitraum.", "Nicht nur die schoenste Marktphase waehlen."],
        ["Deposit / Currency / Leverage", "Kontobedingungen fuer die Simulation.", "Konstant halten, sonst sind Reports nicht vergleichbar."],
        ["Tick Model", "Qualitaet und Geschwindigkeit der Tick-Simulation.", "OHLC fuer schnelle Vorpruefung; realistischere Modi fuer finale Pruefung."],
        ["Visual Mode", "MT5 zeigt Trades im Chart.", "Nur fuer Diagnose und Verstaendnis, nicht fuer Massenlaeufe."],
        ["History & Results", "Profit, Trades, Win Rate, Drawdown.", "Profit nie ohne Tradezahl und Drawdown bewerten."],
    ]
    table(doc, ["Feld", "Bedeutung", "Praxisregel"], rows, [1.55, 2.25, 2.7], "Backtest-Konfiguration")
    h2(doc, "Bedienung: vom EA zum belastbaren Einzelreport")
    numbered(doc, [
        "Expert Advisor waehlen. Der Backtester akzeptiert MT4/MT5-EAs und arbeitet danach mit den Parametern, die der EA selbst bereitstellt.",
        "Symbol und Period setzen. Damit werden Markt und Timeframe an MetaTrader uebergeben; das Tool erfindet keine eigenen Marktdaten.",
        "Datumsfenster bewusst waehlen. Ein Einzeltest sollte einen fachlichen Zweck haben: Funktionspruefung, Plausibilitaet oder Nachtest.",
        "Kontoannahmen konstant halten. Deposit, Currency und Leverage sollten in vergleichbaren Tests nicht staendig wechseln.",
        "Tick Model passend waehlen. Schnelle Vorpruefungen duerfen grober sein, finale Pruefungen brauchen realistischere Datenannahmen.",
        "Start Backtest ausfuehren. Visual Mode nur nutzen, wenn man Trades im Chart beobachten will; Manual Mode nur fuer Diagnose aktivieren.",
        "Report lesen und Historie speichern. Danach nicht nur Profit, sondern Trades, Win Rate, Drawdown, Recovery und Kurvenform bewerten.",
    ])
    action_rows = [
        ["Gen Config", "Startet MetaTrader kurz, um die vom EA gelieferten Inputs als Konfigurationsbasis zu erfassen.", "Nutzen, wenn noch keine Parameterliste vorhanden ist."],
        ["AutoConfig", "Erzeugt sinnvolle Start/Step/Stop-Bereiche aus vorhandenen Werten.", "Als Vorschlag verstehen; fachlich pruefen, bevor optimiert wird."],
        ["Load .set / Save .set", "Liest und schreibt MetaTrader-Parameterdateien.", "SET-Dateien sauber versionieren; nicht mehrere Hypothesen in einer Datei vermischen."],
        ["Start Backtest", "Startet den normalen reproduzierbaren MetaTrader-Lauf.", "Standardaktion fuer Einzeltests."],
        ["Visual Mode", "Oeffnet den MetaTrader sichtbar und zeigt Handelsverlauf im Chart.", "Gut fuer Diagnose, langsam fuer Massenlaeufe."],
        ["Manual Mode", "Laesst MT4/MT5 nach dem Lauf offen.", "Nur verwenden, wenn man Terminalzustand oder Journal pruefen muss."],
        ["Open HTML Report", "Oeffnet den gerenderten Report.", "Immer nutzen, wenn die Tabelle auffaellig gut oder schlecht aussieht."],
        ["Open Directory", "Oeffnet den Reportordner.", "Hilft bei XML/HTML- oder Log-Datei-Kontrolle."],
        ["Delete Selected / Delete All History", "Bereinigt lokale Historie.", "Nur loeschen, wenn klar ist, dass die Ergebnisse nicht mehr gebraucht werden."],
    ]
    table(doc, ["Aktion", "Funktion", "Best Practice"], action_rows, [1.55, 2.35, 2.6], "Backtester-Aktionen")
    param_rows = [
        ["Opt", "Markiert, ob ein Parameter spaeter optimiert werden soll.", "Im Einzeltest meist aus; im Optimizer bewusst aktivieren."],
        ["Variable", "Name des EA-Inputs.", "Nur Parameter aendern, deren Handelslogik verstanden ist."],
        ["Value", "Der feste Wert fuer Einzeltest und SET-Datei.", "Baseline festhalten, bevor Varianten getestet werden."],
        ["Start / Step / Stop", "Optimierungsbereich fuer spaetere Runs.", "Eng genug fuer Rechenzeit, breit genug fuer Plateau-Erkennung."],
    ]
    table(doc, ["Spalte", "Bedeutung", "Praxisregel"], param_rows, [1.35, 2.45, 2.7], "EA-Parameter-Tabelle")
    h2(doc, "Wie man einen Einzeltest liest")
    p(doc, "Ein einzelner Test ist keine Strategie-Freigabe. Er ist ein Messpunkt. Ein hoher Gewinn mit sehr wenigen Trades ist schwach. Ein moderater Gewinn mit vielen Trades, kontrolliertem Drawdown und plausibler Equity-Kurve ist wertvoller. Besonders wichtig sind die ersten und letzten Abschnitte der Equity-Kurve: Oft erkennt man dort, ob der EA nur in einer Marktphase funktioniert hat.")
    image(doc, ROOT / "images" / "backtester_ui3.png", "Einzelreport mit Equity-Kurve, Kennzahlen und Detailstatistik.", width=6.1)
    h2(doc, "Die Grafik lesen: Einzelreport und Strategiequalitaet")
    p(doc, "Der Einzelreport ist die Detailansicht hinter einer Ergebniszeile. Er zeigt typischerweise eine Equity-Kurve, Kennzahlen und Detailstatistik. Die Equity-Kurve beantwortet die wichtigste erste Frage: Wurde das Konto gleichmaessig aufgebaut oder entstand der Gewinn aus wenigen zufaelligen Spruengen? Eine glatte, stetige Kurve mit kontrollierten Rueckgaengen ist besser als eine Kurve, die lange seitwaerts laeuft und nur durch einen einzelnen grossen Trade positiv wird.")
    p(doc, "Die Kennzahlen daneben helfen, die Kurve zu erden. Profit zeigt nur das Endergebnis. Drawdown zeigt das zwischenzeitliche Risiko. Trades zeigen die Stichprobengroesse. Win Rate zeigt, wie oft Trades gewinnen, sagt aber ohne durchschnittlichen Gewinn und Verlust wenig aus. Eine Strategie mit 90 Prozent Win Rate kann schlecht sein, wenn die wenigen Verlusttrades sehr gross sind. Eine Strategie mit 45 Prozent Win Rate kann gut sein, wenn Gewinne groesser als Verluste sind.")
    p(doc, "Wie gut die Strategie im Report ist, darf man deshalb nicht aus einem schoenen Bild allein ableiten. Wenn der Report viele Trades, eine nachvollziehbare Equity-Kurve, niedrige bis moderate Drawdowns und eine stabile Leistung ueber verschiedene Marktphasen zeigt, ist das ein gutes erstes Signal. Wenn nur wenige Trades vorhanden sind oder ein grosser Einzeltrade das Ergebnis rettet, ist die Aussage schwach. Der Report ist also ein Pruefstein, aber noch kein endgueltiges Urteil.")
    callout(
        doc,
        "Warum Einzeltests gebraucht werden",
        "Einzeltests schuetzen vor teurer Blindoptimierung. Wenn ein EA schon im Basistest keine Trades ausfuehrt, Reports nicht schreibt oder nur auf einem Zufallsereignis gewinnt, muss man nicht stundenlang optimieren.",
        fill=PALE_GREEN,
        accent=GREEN,
    )


def add_multibacktester_section(doc: Document, graphics: dict[str, Path]):
    h1(doc, "5. Multi-Backtester: Maerkte und Timeframes vergleichen")
    image(doc, graphics["market_timeframes"], "Maerkte und Timeframes als Matrix: Der Multi-Backtester macht systematische Vergleiche moeglich.", width=6.45)
    h2(doc, "Warum dieser Bereich gebraucht wird")
    p(doc, "Eine Strategie ist selten auf jedem Markt und jedem Timeframe gleich gut. Der Multi-Backtester testet einen EA mit denselben Konto- und Zeitraumannahmen ueber viele Symbol/Perioden-Kombinationen. Dadurch erkennt man, ob eine Strategie nur zufaellig auf einem Chart funktioniert oder ob es ein wiederkehrendes Muster gibt.")
    p(doc, "Auch hier wird keine neue Simulationslogik gebaut. Jeder einzelne Lauf ist ein normaler MetaTrader-Backtest. Das Projekt automatisiert nur die Serienausfuehrung und die Auswertung. Genau diese Gleichheit ist wichtig: Wenn alle Runs mit denselben Annahmen laufen, sind die Ergebnisse fair vergleichbar.")
    p(doc, "Der eigentliche Nutzen ist die Frage nach der Uebertragbarkeit: Kann eine bestehende Strategie nur auf einem einzigen Symbol gut aussehen, oder funktioniert ihre Logik auch auf mehreren Datenreihen? Wenn ein EA auf EURUSD, GBPUSD und USDJPY in verwandten Timeframes brauchbar performt, ist das ein besseres Zeichen als ein isolierter Spitzenlauf. Mehrere Maerkte bedeuten mehrere Preisreihen, andere Volatilitaeten, andere Sessions und andere Trendphasen. Eine Strategie, die dort nicht sofort zusammenbricht, ist weniger verdaechtig, nur auf eine historische Kurve ueberoptimiert zu sein.")
    callout(
        doc,
        "Einfaches Beispiel",
        "Stellen Sie sich eine Breakout-Strategie vor. Auf EURUSD H1 gewinnt sie 18 Prozent, auf GBPUSD H1 15 Prozent und auf USDJPY H1 11 Prozent bei aehnlichem Drawdown. Das ist kein Beweis, aber ein starkes Plausibilitaetssignal. Wenn dieselbe Strategie dagegen nur auf EURUSD M15 gewinnt und auf allen anderen Maerkten verliert, sollte man sehr misstrauisch sein.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    h2(doc, "Was ein Markt im Forex-Kontext ist")
    p(doc, "Im Forex-Handel ist ein Markt meist ein Waehrungspaar. EURUSD bedeutet: Der Euro ist die Basiswaehrung, der US-Dollar die Kurswaehrung. Der Kurs zeigt, wie viele US-Dollar fuer einen Euro bezahlt werden. Investopedia beschreibt Currency Pairs genau als diesen Vergleich zwischen Basis- und Kurswaehrung [S6].")
    p(doc, "Majors wie EURUSD, GBPUSD oder USDJPY sind besonders liquide. Crosses wie EURGBP oder EURJPY enthalten nicht direkt den US-Dollar und koennen andere Bewegungsmuster zeigen. Exotischere Paare oder Rohstoffsymbole wie XAUUSD und XTIUSD koennen hoehere Volatilitaet, breitere Spreads oder staerkere Session-Effekte haben. Genau deshalb sollte ein EA nicht blind auf alle Symbole optimiert werden.")
    h2(doc, "Was Timeframes bedeuten")
    p(doc, "Ein Timeframe ist die Zeiteinheit der Kerzen: M1 ist eine Minute, H1 eine Stunde, D1 ein Tag. Kleine Timeframes liefern mehr Signale, aber auch mehr Rauschen und hoehere Anforderungen an Tickdaten, Spread und Ausfuehrung. Groessere Timeframes liefern weniger Trades, aber oft stabilere Marktstrukturen. Multiple-Timeframe-Analyse nutzt kurze, mittlere und lange Betrachtungen, um Trend, Einstieg und Kontext zu trennen [S7].")
    p(doc, "Der Multi-Backtester hilft dabei, einen passenden Arbeitsbereich zu finden. Wenn eine Strategie auf M1 nur durch Rauschen gewinnt, auf H1 aber stabilere Ergebnisse zeigt, ist das ein wichtiger Hinweis. Umgekehrt kann eine Scalping-Strategie auf D1 schlicht zu wenig Trades erzeugen.")
    image(doc, ROOT / "images" / "multi-backtester-config.png", "Multi-Backtester-Konfiguration: ein EA, ein Zeitraum, viele Symbol- und Timeframe-Jobs.", width=6.3)
    p(doc, "Die Konfigurationsgrafik zeigt den Multi-Backtester als Experiment-Planer. Oben werden die gemeinsamen Annahmen gesetzt: EA, Zeitraum, Konto- und Testerbedingungen. Darunter werden die Maerkte und Timeframes markiert, die mit genau diesen Annahmen getestet werden sollen. Das ist wichtig, weil nur gleich konfigurierte Runs fair verglichen werden koennen. Wenn EURUSD mit anderem Zeitraum als USDJPY getestet wird, vergleicht man nicht mehr die Strategie, sondern unterschiedliche Experimente.")
    p(doc, "Der Nutzer sieht hier also keine einzelne Strategieentscheidung, sondern den Versuchsaufbau: Welche Datenreihen werden gegen dieselbe Handelslogik gehalten? Genau diese Sicht macht Ueberoptimierung sichtbarer. Eine Strategie, die nur auf einem einzigen Kreuzchen gewinnt, ist schwach. Eine Strategie, die in mehreren benachbarten Feldern der Matrix brauchbar bleibt, verdient eine genauere Pruefung.")
    h2(doc, "Praktische Bedienung")
    rows = [
        ["Expert Advisor", "Ein EA fuer alle Jobs.", "Parameter muessen fuer alle ausgewaehlten Maerkte plausibel sein."],
        ["Dates", "Ein globaler Zeitraum.", "Nur gleiche Zeitfenster machen Ergebnisse vergleichbar."],
        ["Symbols", "Checkbox-Liste plus Add Custom.", "Custom Symbols zuerst in MT5 verfuegbar machen."],
        ["Timeframes", "M1 bis MN1.", "Nicht alles blind markieren; Hypothese vorher formulieren."],
        ["Presets", "EA, Symbolauswahl, Timeframes und Parameter-Snapshot.", "Fuer wiederholbare Markt-Screenings nutzen."],
        ["Open Multi-Report", "Aggregierter HTML-Report.", "Nach Clustern suchen: Marktgruppe, Timeframe, Tradezahl, Drawdown."],
    ]
    table(doc, ["Element", "Funktion", "Warum wichtig"], rows, [1.35, 2.2, 2.95], "Multi-Backtester-Referenz")
    h2(doc, "Wie man Markt- und Timeframe-Ergebnisse interpretiert")
    p(doc, "Die Ergebnistabelle ist keine Rangliste, aus der man automatisch die oberste Zeile nimmt. Sie ist eine Karte der Strategieeigenschaften. Wenn ein EA auf EURUSD, GBPUSD und USDJPY in H1/H4 solide läuft, aber auf M1/M5 stark schwankt, spricht das eher fuer eine mittelfristige Idee. Wenn nur XAUUSD M15 auffaellig gut ist, muss man sehr streng pruefen, ob dort ein echter Effekt oder nur ein Sonderfall vorliegt.")
    p(doc, "Achten Sie besonders auf Cluster. Ein Cluster bedeutet: Mehrere verwandte Maerkte oder mehrere benachbarte Timeframes zeigen aehnliche Qualitaet. Genau das passt zum Plateau-Denken. Einzelne Ausreisser sind dagegen verdachtig, selbst wenn sie spektakulaer aussehen.")
    result_rows = [
        ["Robot / Symbol / Period", "Identifiziert jeden Einzelrun.", "Nach Gruppen sortieren: erst Marktgruppe, dann Timeframe."],
        ["Trades", "Anzahl der abgeschlossenen Trades.", "Zu wenige Trades liefern keine robuste Aussage."],
        ["Win Rate", "Anteil gewonnener Trades.", "Nur mit Payoff und Drawdown sinnvoll interpretieren."],
        ["Drawdown", "Maximaler Rueckgang.", "Hohe Profite mit sehr hohem Drawdown sind selten professionell handelbar."],
        ["Recovery Factor", "Gewinn im Verhaeltnis zum Drawdown.", "Guter Schnellindikator fuer Risiko-Nutzen-Verhaeltnis."],
        ["Profit / Status", "Endergebnis und technischer Status.", "Fehlerhafte Runs zuerst technisch klaeren, nicht fachlich bewerten."],
    ]
    table(doc, ["Spalte", "Bedeutung", "Leseregel"], result_rows, [1.55, 2.25, 2.7], "Ergebnisliste des Multi-Backtesters")
    image(doc, ROOT / "images" / "multi-backtester-results.png", "Multi-Backtester-Ergebnisse: Vergleich vieler Einzelruns mit Profit, Trades, Drawdown und Status.", width=6.3)
    p(doc, "Die Ergebnisgrafik ist die Landkarte nach dem Experiment. Jede Zeile ist ein normaler MetaTrader-Backtest, aber die Tabelle macht die vielen Einzelreports vergleichbar. Man sucht nicht nur die hoechste Profitzahl, sondern Muster: Welche Symbole bleiben positiv? Welche Timeframes brechen ein? Gibt es genug Trades? Ist der Drawdown in allen guten Zeilen vergleichbar niedrig oder gibt es einzelne Ausreisser?")
    p(doc, "Besonders wertvoll sind zusammenhaengende Ergebnisgruppen. Wenn H1 und H4 auf mehreren Major-Paaren solide sind, spricht das fuer eine robuste Zeitebene. Wenn nur ein exotisches Symbol mit wenigen Trades sehr gut aussieht, ist Vorsicht angebracht. Der Multi-Backtester ist damit ein Fruehwarnsystem gegen Ein-Markt-Optimierung.")
    h2(doc, "Multi-Backtester-Aktionen: warum dieser Werkzeugkasten wichtig ist")
    p(doc, "Die Aktionen im Multi-Backtester sind nicht nur Bedienknöpfe. Sie steuern den Marktvergleich als Experiment. Start Batch erzeugt die Datenbasis, Presets machen das Experiment wiederholbar, Open Multi-Report zeigt das Gesamtbild, und Show Single Report erlaubt die Detailprüfung auffälliger Runs. So wird aus vielen Einzeltests eine strukturierte Antwort auf die Frage: Ist diese Strategie generell brauchbar oder nur auf einem Chart schön?")
    action_rows = [
        ["Start Batch", "Startet alle markierten Symbol/Timeframe-Kombinationen nacheinander.", "Vorher kontrollieren, wie viele Jobs entstehen."],
        ["Cancel", "Bricht die laufende Batch-Sequenz ab.", "Nutzen, wenn ein Symbol haengt oder die Annahmen falsch waren."],
        ["Presets: Neu/Speichern/Aendern/Loeschen", "Speichert wiederkehrende EA-, Markt- und Timeframe-Sets.", "Fuer vergleichbare Screenings feste Presets verwenden."],
        ["Add Custom", "Fuegt ein nicht vorgegebenes Symbol hinzu.", "Nur wenn das Symbol in MetaTrader existiert oder importiert wurde."],
        ["Open Multi-Report Node", "Oeffnet den aggregierten HTML-Report.", "Erster Ort fuer Clusteranalyse und Portfolio-Blick."],
        ["Show Single Report", "Oeffnet den Report eines markierten Einzelruns.", "Bei Ausreissern immer den Einzelreport lesen."],
        ["Delete Batch / Delete Selected Runs", "Bereinigt Batch-Historie.", "Nur nach Export oder bewusster Verwerfung nutzen."],
    ]
    table(doc, ["Aktion", "Funktion", "Best Practice"], action_rows, [1.75, 2.35, 2.4], "Multi-Backtester-Aktionen")
    callout(
        doc,
        "Warnung vor Auswahlbias",
        "Wenn man 50 Kombinationen testet und nur den besten Run betrachtet, hat man fast sicher einen Glueckstreffer erzeugt. Der Multi-Backtester ist ein Screening-Werkzeug, keine Freigabeinstanz.",
        fill="FDECEC",
        accent=RED,
    )
    source_note(doc, "Quellen: Forex-Paare und Marktzeiten nach Investopedia [S6, S8], Multiple-Timeframe-Analyse nach Investopedia [S7].")


def add_optimizer_section(doc: Document, graphics: dict[str, Path]):
    h1(doc, "6. Optimizer: Optimierung, Plateau-Denken und robuste Auswahl")
    image(doc, ROOT / "images" / "backtester_optimizer.png", "Optimizer: Suchraum, Parameter, Combined Analysis und Ergebnisliste.", width=6.35)
    h2(doc, "Die Grafik lesen: vom Suchraum zur Auswahl")
    p(doc, "Der Optimizer-Screenshot zeigt die eigentliche Strategie-Werkstatt. Im oberen Bereich werden EA, Symbol, Timeframe, Zeitraum, Kontoannahmen, Tickmodell, Optimierungsmodus und Forward-Einstellung gesetzt. In der Parameter-Tabelle wird dann entschieden, welche EA-Inputs optimiert werden und welche Werte als Start, Step und Stop in den Suchraum gehen.")
    p(doc, "Der untere Bereich ist die Auswertungszone. Dort werden Optimierungspasses, Forward-Ergebnisse und Combined-Analysis-Daten sichtbar. Genau hier unterscheidet sich das Projekt stark von einem reinen MetaTrader-Lauf: Die Ergebnisse werden nicht nur gesammelt, sondern gefiltert, gewichtet, zusammengefuehrt und spaeter durch Sensitivitaet, KI und OOS-Test weiter geprueft.")
    callout(
        doc,
        "Was der Screenshot fachlich bedeutet",
        "Der Optimizer ist nicht dafuer da, den hoechsten historischen Profit zu finden. Er soll zeigen, ob ein EA in einem sinnvollen Parameterraum viele brauchbare Kandidaten erzeugt, ob Forward-Daten die Auswahl bestaetigen und ob die guten Werte in stabilen Zonen liegen. Darum gehoeren Suchraum, Filter, Score und Forward-Split immer zusammen.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    h2(doc, "Was Optimierung bedeutet")
    p(doc, "Optimierung bedeutet, dass ein Expert Advisor nicht nur einmal mit festen Parametern getestet wird, sondern viele Male mit unterschiedlichen Eingabewerten. MetaTrader kann diese Kombinationen komplett oder genetisch durchsuchen. Ziel ist nicht, irgendeinen historischen Maximalwert zu finden, sondern Parameter zu finden, die plausibel, stabil und handelbar sind.")
    p(doc, "MetaTrader stellt dafuer die Strategy-Tester-Optimierung bereit. Das Backtester-Projekt nutzt diese vorhandene Funktion, fuegt aber Auswertungsebenen hinzu, die in MetaTrader vermisst werden: Combined Analysis, Filtersettings, Score-Gewichtung, Forward-Konsistenz, Advanced Evaluator, Sensitivitaetsanalyse, KI-Bewertung und Step-7-Validierung.")
    p(doc, "Warum optimiert man ueberhaupt? Maerkte bleiben nicht gleich. Volatilitaet, Spread, Trendlaenge, Handelszeiten und Reaktionsgeschwindigkeit veraendern sich. Ein Parameter-Set, das vor zwei Jahren perfekt war, kann heute zu eng, zu langsam oder zu aggressiv sein. Optimierung ist deshalb keine Spielerei, sondern eine kontrollierte Methode, um eine Strategie an Marktbedingungen anzupassen und gleichzeitig zu pruefen, ob sie ueberhaupt eine robuste Grundidee besitzt.")
    callout(
        doc,
        "Einfaches Beispiel: gleitender Durchschnitt",
        "Ein EA nutzt einen gleitenden Durchschnitt als Trendfilter. Mit Periode 20 reagiert er schnell, produziert aber viele Fehlsignale. Mit Periode 80 reagiert er langsam und verpasst Einstiege. Die Optimierung testet Werte wie 20, 30, 40, 50, 60, 70 und 80. Interessant ist nicht nur, welcher Wert den hoechsten Profit hatte, sondern ob eine ganze Zone, zum Beispiel 40 bis 60, brauchbar bleibt.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    p(doc, "Schon die Verteilung der Optimierungsergebnisse sagt viel ueber die Strategie. Wenn 80 Prozent aller getesteten Parameterkombinationen Verlust machen und nur ein einzelner Pass extrem gut aussieht, ist die Grundidee wahrscheinlich schwach oder stark ueberangepasst. Wenn dagegen viele Kombinationen zumindest stabil oder leicht positiv sind, spricht das fuer eine robustere Logik. Der Optimizer ist also nicht nur eine Suchmaschine fuer den besten Wert, sondern auch ein Diagnosewerkzeug fuer die Qualitaet der Strategieidee.")
    source_note(doc, "MetaQuotes beschreibt Optimierung als wiederholtes Testen desselben Trading-Roboters mit unterschiedlichen Parametern und weist auf genetische Algorithmen zur Reduktion grosser Kombinationen hin [S2].")
    settings_rows = [
        ["Expert Advisor", "Der zu optimierende EA.", "Nur EAs optimieren, die im Backtest technisch sauber laufen."],
        ["Symbol / Period", "Markt und Timeframe fuer die Optimierung.", "Aus Multi-Backtester-Screening ableiten, nicht blind waehlen."],
        ["Date Range", "In-Sample-Zeitraum plus optionaler Forward-Split.", "OOS-Fenster vorab reservieren und nicht in der Auswahl verbrauchen."],
        ["Deposit / Currency / Leverage", "Kontobasis wie im MetaTrader.", "Konstant halten, sonst sind Optimierungen nicht vergleichbar."],
        ["Tick Model", "Modell der historischen Tick-/Kursdaten.", "Finale Kandidaten mit hoeherer Datenqualitaet pruefen."],
        ["Opt. Mode", "Slow Complete oder Fast Genetic Algorithm.", "Complete fuer kleine Suchraeume, Genetic fuer grosse Raeume."],
        ["Opt. Criterion", "Zielgroesse der MetaTrader-Optimierung.", "Recovery oder Sharpe oft robuster als reiner Maximalprofit."],
        ["Forward Test / Forward Date", "Automatischer oder eigener Split in neue Daten.", "Immer aktivieren, wenn robuste Auswahl das Ziel ist."],
    ]
    table(doc, ["Einstellung", "Funktion", "Professionelle Nutzung"], settings_rows, [1.55, 2.35, 2.6], "Optimization Settings")

    h2(doc, "Warum Spitzen gefaehrlich sind")
    p(doc, "Eine Optimierung kann ein einzelnes spektakulaeres Ergebnis liefern: maximaler Profit, fast perfekte Equity-Kurve, geringer Drawdown. Wenn dieser Punkt aber direkt neben schlechten Nachbarwerten liegt, ist er fachlich schwach. Schon eine kleine Veraenderung von Stop-Loss, Take-Profit oder Filterperiode wuerde das Ergebnis zerstoeren. Genau das ist typisch fuer Ueberoptimierung.")
    p(doc, "QuantStart nennt Optimierungsbias beziehungsweise Curve Fitting eine der gefaehrlichsten Backtest-Fallen und empfiehlt unter anderem Sensitivitaetsanalyse: Parameter inkrementell veraendern und auf glatte Performance-Oberflaechen achten [S3]. Neuere Forschung spricht explizit von Parameter-Plateaus als stabilen Regionen, die Overfitting-Risiken reduzieren koennen [S4].")
    image(doc, graphics["plateau_optimization"], "Plateau-Optimierung: gesucht wird eine stabile Zone, nicht ein einzelner historischer Spitzenpunkt.", width=6.45)

    h2(doc, "Was mit Plateau- oder Tableau-Optimierung gemeint ist")
    p(doc, "Im Alltag wird manchmal von Tableau-Optimierung gesprochen; fachlich ist hier Plateau-Optimierung gemeint. Ein Plateau ist ein Bereich im Parameterraum, in dem mehrere benachbarte Werte aehnlich gute Ergebnisse liefern. Das ist ein starkes Zeichen: Die Strategie haengt dann nicht an exakt einem Wert, sondern an einer robusteren Marktlogik.")
    p(doc, "Beispiel: Wenn TakeProfit 48, 49, 50, 51 und 52 Punkte alle solide Ergebnisse liefern, wirkt die Zone stabil. Wenn nur 50 Punkte hervorragend ist und 49 sowie 51 sofort verlieren, ist das ein Peak. Peaks sehen in der Vergangenheit attraktiv aus, brechen aber in neuen Daten haeufig ein.")
    h3(doc, "Wie das Projekt Plateaus sucht")
    numbered(doc, [
        "Combined Analysis verbindet Backtest- und Forward-Ergebnisse, damit reine In-Sample-Gewinner nicht dominieren.",
        "Filtersettings entfernen Kandidaten mit zu wenig Trades, zu hohem Drawdown oder schlechtem Forward-Verhalten.",
        "Score-Gewichtung belohnt nicht nur Profit, sondern auch Konsistenz, Sample Size, Forward-Trades und Recovery.",
        "SensitivityRunner verschiebt Parameter um ihren optimierten Wert und misst CV-Werte sowie Kurvenform.",
        "KI-Analyse beschreibt, ob Kurven wie Plateau, Peak oder Klippe aussehen.",
        "Step 7 testet finale Kandidaten auf unberuehrten Daten.",
    ])

    h2(doc, "Filtersettings richtig verstehen")
    rows = [
        ["BT/FW Profit", "Mindestgewinn in Backtest und Forward.", "Forward-Verlierer nicht als robust betrachten."],
        ["BT/FW Trades", "Mindestzahl Trades.", "Schuetzt vor Glueckstreffern mit kleiner Stichprobe."],
        ["BT/FW Drawdown", "Maximal tolerierter Rueckgang.", "Profit ohne Risiko ist keine robuste Kennzahl."],
        ["Expected Payoff", "Durchschnitt pro Trade.", "Sehr kleine Payoffs sind kosten- und slippage-anfaellig."],
        ["Sharpe / Recovery", "Risikoadjustierte Stabilitaet.", "Hilft, glattere Strategien hoeher zu bewerten."],
        ["Consistency", "Forward Profit relativ zu Backtest Profit.", "Schuetzt vor Kandidaten, die nach dem Split einbrechen."],
    ]
    table(doc, ["Filter", "Bedeutung", "Sinn"], rows, [1.45, 2.25, 2.8], "Wichtige Filter im Optimizer")
    p(doc, "Der Filterdialog wird gebraucht, weil Optimierungen sehr viele Ergebnisse erzeugen. Ohne Filter schaut man schnell nur auf Profit und uebersieht, dass ein Kandidat zu wenige Trades, zu hohen Drawdown oder ein schlechtes Forward-Ergebnis hat. Der Filterdialog ist wie ein Sieb: Er entfernt Ergebnisse, die zwar spektakulaer aussehen koennen, aber gegen Mindestanforderungen verstossen.")
    p(doc, "Ein einfaches Beispiel: Eine Optimierung liefert 3.000 Passes. Der beste Pass macht 4.000 Euro Gewinn, hat aber nur 6 Trades. Ein anderer Pass macht 2.200 Euro Gewinn, hat 180 Trades, moderaten Drawdown und bleibt im Forward positiv. Ohne Filter wuerde man vielleicht den ersten Kandidaten waehlen. Mit Filterdialog fallen zu kleine Stichproben heraus, und die realistischeren Kandidaten werden sichtbar.")
    detail_rows = [
        ["Filter aktiv", "Schaltet die Combined-Analysis-Filter ein.", "Bei ersten Rohdaten aus lassen, danach bewusst aktivieren."],
        ["Nur Passes mit Forward-Ergebnis", "Entfernt reine Backtest-Passes aus der Combined-Auswahl.", "Für robuste Optimierung normalerweise eingeschaltet lassen."],
        ["Min BT Profit / Min FW Profit", "Mindestprofit in In-Sample und Forward.", "FW Profit nie ignorieren, sonst gewinnt der historische Peak."],
        ["Min BT Trades / Min FW Trades", "Mindestanzahl an Trades.", "Default im Code: 100 Backtest-Trades und 15 Forward-Trades als Startpunkt."],
        ["Max BT DD / Max FW DD", "Obergrenze fuer Drawdown in Prozent.", "Nicht als kosmetischen Filter behandeln; Drawdown ist Ueberlebensrisiko."],
        ["Min Payoff", "Mindestgewinn je Trade.", "Schuetzt vor Strategien, die durch Kosten/Slippage verschwinden."],
        ["Min Sharpe", "Mindestwert fuer risikoadjustierte Rendite.", "Hilft gegen wilde Equity-Kurven."],
        ["Min Recovery", "Mindestverhaeltnis aus Gewinn und Drawdown.", "Guter Filter fuer Risiko-Nutzen-Verhaeltnis."],
        ["Mindest-Score", "Score-Schwelle 0-100 mit Presets Low/Med/High.", "Nicht zu hoch starten; zuerst verstehen, warum Kandidaten ausscheiden."],
        ["Mindest-Konsistenz", "Forward/Backtest-Verhaeltnis, Code-Presets 0.4/0.6/0.8.", "Je strenger, desto weniger, aber belastbarere Kandidaten."],
    ]
    table(doc, ["Filterdialog", "Was er prueft", "Best Practice"], detail_rows, [1.55, 2.35, 2.6], "Filterdialog im Detail")
    p(doc, "Score-Gewichtung wird gebraucht, weil es nicht die eine perfekte Kennzahl gibt. Profit allein belohnt oft riskante Kandidaten. Drawdown allein kann zu defensiv sein. Tradezahl allein sagt nichts ueber Qualitaet. Die Score-Gewichtung verbindet mehrere Ranking-Kriterien zu einer sortierbaren Bewertung. So koennen die besten Strategien herausgepickt werden, ohne dass ein einzelner Messwert das Ergebnis dominiert.")
    p(doc, "In der Praxis bedeutet das: Wer robuste Portfolio-Kandidaten sucht, kann Forward-Profit, Konsistenz, Recovery und Sample Size hoeher gewichten. Wer nur Ideen schnell vorsortiert, kann die Profit-Komponenten etwas staerker gewichten. Wichtig ist, dass die Gewichtung zur Frage passt. Das Ranking ist kein Urteil des Marktes, sondern eine bewusst gesetzte Bewertungslogik.")
    score_rows = [
        ["BT Profit", "Profitabilitaet im Backtest.", "Wichtig, aber allein gefaehrlich."],
        ["FW Profit", "Profitabilitaet im Forward-Fenster.", "Soll hohes Gewicht haben, weil neue Daten entscheidend sind."],
        ["Konsistenz", "Verhaeltnis Forward zu Backtest.", "Schuetzt vor In-Sample-Wundern."],
        ["Risiko-Verhaeltnis", "Drawdown- und Risikobeitrag.", "Hoher Profit mit hohem DD wird abgewertet."],
        ["Sharpe / Equity Consistency", "Glattheit und risikoadjustierter Verlauf.", "Belohnt weniger sprunghafte Strategien."],
        ["Sample Size", "Stichprobengroesse.", "Viele Trades geben mehr Aussagekraft als wenige Treffer."],
        ["FW Trades", "Trades im Forward.", "Ohne Forward-Trades ist Forward-Profit kaum belastbar."],
        ["Recovery", "Erholung nach Rueckgaengen.", "Zeigt, ob Gewinn und Risiko in einem sinnvollen Verhaeltnis stehen."],
    ]
    table(doc, ["Score-Saeule", "Bewertet", "Warum sie gebraucht wird"], score_rows, [1.55, 2.25, 2.7], "Score-Gewichtung")
    image(doc, ROOT / "images" / "backtester_score_weighting.png", "Score-Gewichtung: Welche Eigenschaften im Ranking zaehlen sollen.", width=4.9)
    p(doc, "Die Score-Grafik zeigt die Gewichtungslogik hinter dem Ranking. Jede Saeule steht fuer eine Eigenschaft, die in die Gesamtbewertung einfliessen kann: Profit, Forward-Leistung, Konsistenz, Risiko, Sample Size, Recovery und weitere Stabilitaetsmerkmale. Durch die Gewichtung wird festgelegt, welche Eigenschaften im konkreten Auswahlprozess wichtiger sind.")
    p(doc, "Das ist noetig, weil verschiedene Strategietypen unterschiedliche Profile haben. Eine kurzfristige Scalping-Strategie braucht viele Trades und geringe Kostenempfindlichkeit. Eine langsamere H4-Strategie darf weniger Trades haben, muss dafuer aber sauberere Drawdowns und bessere Konsistenz zeigen. Die Score-Gewichtung macht diese Prioritaeten explizit, statt sie unbewusst in die Auswahl einzuschleusen.")
    image(doc, ROOT / "images" / "backtester_advanced_evaluator.png", "Advanced Evaluator: Detailanalyse eines Passes statt blindem Vertrauen in eine Tabellenzeile.", width=6.2)
    p(doc, "Die Advanced-Evaluator-Grafik ist die Lupenansicht auf einen einzelnen Kandidaten. Oben sieht man typischerweise die zentralen Kennzahlen und Bewertungsfelder, unten oder daneben die Detailkurven, Parameter und Erklaerungen. Damit wird die abstrakte Pass-Zeile wieder zu einer konkreten Strategie: Man sieht, wodurch der Score entstanden ist und ob die Kennzahlen zusammen ein stimmiges Bild ergeben.")
    p(doc, "Der Advanced Evaluator ist eine Auswahlhilfe fuer die besten Strategien. Er nimmt einen Pass aus der Tabelle und macht ihn verstaendlich: Wie gut war der Backtest? Wie stark war der Forward? Wie sieht die Equity-Kurve aus? Passen Profit, Drawdown, Recovery, Tradezahl und Parameter zusammen? Damit wird aus einer Tabellenzeile eine Strategieakte. Gerade wenn die Top 20 sehr aehnlich aussehen, hilft der Evaluator, nicht nur den hoechsten Score zu nehmen, sondern den plausibelsten Kandidaten.")
    callout(
        doc,
        "Auswahlhilfe statt Autopilot",
        "Der Advanced Evaluator entscheidet nicht automatisch, welche Strategie live gehen darf. Er hilft, die besten Kandidaten zu verstehen und schlechte Kandidaten zu entlarven: zu wenige Trades, Forward-Einbruch, unschoene Equity-Kurve, zu hoher Drawdown oder ein Parameter-Set, das fachlich keinen Sinn ergibt.",
        fill=LIGHT_GRAY,
        accent=BLUE,
    )
    image(doc, ROOT / "images" / "backtester_consistency_ratio.png", "Konsistenzanalyse: Forward- und Backtest-Leistung werden ins Verhaeltnis gesetzt.", width=5.7)
    p(doc, "Die Konsistenzanalyse erklaert, wie Backtest- und Forward-Leistung zusammenhaengen. Genau solche Erklaerungen erscheinen in der Anwendung oft ueber die kleinen Info-Buttons. Wenn man auf das i klickt, oeffnet sich eine fachliche Kurz-Doku direkt an der Stelle, an der sie gebraucht wird. Zusaetzlich gibt es im Manual-Tab eine breitere Dokumentation. Die Grafik zur Konsistenz macht sichtbar, ob der Forward noch im Verhaeltnis zum Backtest steht oder ob der Kandidat nach dem Split zusammenfaellt.")
    p(doc, "Ein Beispiel: Im Backtest macht ein Pass 10.000 Euro, im Forward aber nur 200 Euro oder sogar Verlust. Dann war die historische Leistung wahrscheinlich nicht stabil. Macht der Backtest 4.000 Euro und der Forward 1.800 Euro bei genug Trades, ist das weniger spektakulaer, aber oft glaubwuerdiger. Konsistenz ist deshalb ein Anti-Curvefitting-Werkzeug.")
    image(doc, ROOT / "images" / "backtester_sensitivity.png", "Sensitivitaetsdetails: CV-Werte und Kurven zeigen Peaks, Plateaus und Klippen.", width=6.35)
    p(doc, "Die Sensitivitaetsanalyse wird gemacht, um zu sehen, wie empfindlich eine Strategie auf kleine Parameterveraenderungen reagiert. Eine robuste Strategie darf nicht davon abhaengen, dass ein Wert exakt 50 und nicht 49 oder 51 ist. Wenn kleine Aenderungen sofort grosse Verluste erzeugen, ist die Strategie fragil. Wenn mehrere Nachbarwerte aehnlich gut bleiben, spricht das fuer ein Plateau.")
    p(doc, "In der Grafik sieht man pro Parameter Kennlinien und CV-Werte. Die Kennlinie zeigt, wie sich Profit oder Score veraendert, wenn der Parameter um den optimierten Wert herum verschoben wird. Der CV-Wert fasst zusammen, wie stark die Ergebnisse streuen. Daneben erklaert die Anwendung, ob das Verhalten eher robust, akzeptabel oder fragil wirkt. Genau hier lernt der Nutzer, ob die Optimierung eine breite stabile Zone gefunden hat oder nur einen spitzen historischen Treffer.")
    callout(
        doc,
        "Merksatz fuer Optimierung",
        "Nicht der hoechste historische Pass ist das Ziel. Das Ziel ist eine robuste Parameterzone, die auch im Forward, in der Sensitivitaet und spaeter im OOS-Test plausibel bleibt.",
        fill=PALE_GREEN,
        accent=GREEN,
    )


def add_workflow_section(doc: Document):
    h1(doc, "7. Workflow, Robustheit, KI und Controlling")
    image(doc, ASSET_DIR / "workflow.png", "Acht Arbeitsphasen: Vorbereitung plus die sieben UI-Schritte des Workflow Automators.", width=6.45)
    p(doc, "Die Workflow-Grafik ordnet die Arbeit in eine Prozesskette. Links beginnt die Vorbereitung: Daten, Settings, EA und OOS-Plan muessen stehen, bevor Optimierung sinnvoll ist. Danach folgen die UI-Schritte des Workflow Automators: Auswahl, Optimierung, Filter/Diversitaet, Sensitivitaet, KI-Bewertung, Portfolio-Export und echte Out-of-Sample-Validierung.")
    p(doc, "Wichtig ist die Leserichtung: Jede Stufe reduziert Unsicherheit, ersetzt aber nicht die naechste. Eine gute Optimierung ersetzt keine Sensitivitaetsanalyse. Eine gute KI-Bewertung ersetzt keinen OOS-Test. Ein Export ersetzt kein Controlling. Die Grafik macht deshalb sichtbar, warum robuste Strategieentwicklung ein Prozess ist und kein einzelner Klick.")
    h2(doc, "Warum der Workflow gebraucht wird")
    p(doc, "Ein einzelner guter Report ist keine belastbare Strategieentscheidung. Der Workflow Automator macht aus vielen Einzelschritten eine gefuehrte Pipeline. Dadurch sieht der Nutzer, ob eine Strategie erst technisch getestet, bereits optimiert, schon diversifiziert, sensitivitaetsgeprueft, KI-bewertet, exportiert oder wirklich out-of-sample validiert wurde.")
    p(doc, "Der Workflow spart vor allem Arbeit und verhindert Auslassungen. Frueher musste man jeden Schritt von Hand machen: Optimierung starten, Ergebnisse exportieren, Kandidaten sortieren, Filter anwenden, Sensitivitaet testen, KI-Bericht erzeugen, Strategien in Ordner kopieren und spaeter ein echtes OOS-Fenster nachtesten. Bei wenigen Strategien ist das machbar. Bei grossen Mengen wird es langsam, fehleranfaellig und unuebersichtlich.")
    p(doc, "Der Workflow automatisiert diese Optimierungsschritte. Er ist deshalb besonders wertvoll, wenn viele Strategien oder viele EAs systematisch optimiert werden sollen. Der Nutzer muss nicht mehr jede Zwischenentscheidung aus dem Gedaechtnis rekonstruieren. Das Programm fuehrt durch die Pipeline, merkt sich den Zustand und macht sichtbar, welche Stufe bereits erledigt ist.")
    callout(
        doc,
        "Einfaches Beispiel: von Hand vs. Workflow",
        "Ohne Workflow optimiert man montags einen EA, filtert dienstags Kandidaten, vergisst mittwochs, welche Forward-Einstellung genutzt wurde, und exportiert freitags versehentlich einen nicht validierten Pass. Mit Workflow bleibt die Kette zusammen: Setup, Optimierung, Filter, Sensitivitaet, KI, Portfolio und Step-7-Validierung sind als Prozess nachvollziehbar.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    phases = [
        ["0 Vorbereitung", "Settings, Daten, EA und OOS-Plan klaeren.", "Ohne saubere Ausgangslage sind alle Ergebnisse schwach."],
        ["1 Strategie-Auswahl", "EA, Symbol, Zeitraum, Konto, Suchraum.", "Nur verstandene Parameter optimieren."],
        ["2 Optimierung", "Algorithmus, Ziel, Forward-Modus.", "Forward ist Pflicht fuer robuste Vorauswahl."],
        ["3 Filter & Diversitaet", "Profit, Trades, Drawdown, Unterschiedlichkeit.", "Nicht fuenf fast identische Paesse behalten."],
        ["4 Sensitivitaet", "Parameter-Sweeps und CV.", "Plateaus suchen, Klippen vermeiden."],
        ["5 KI-Bewertung", "Kurvenform und Stabilitaet beschreiben lassen.", "KI erklaert, entscheidet aber nicht allein."],
        ["6 Portfolio", "3-5 finale Kandidaten exportieren.", "Export ist noch keine Live-Freigabe."],
        ["7 OOS-Validierung", "Spaeteres unberuehrtes Fenster testen.", "Nur bestandene Kandidaten gehoeren in den Best-Ordner."],
    ]
    table(doc, ["Phase", "Aufgabe", "Warum wichtig"], phases, [1.45, 2.55, 2.5], "Die acht Phasen im praktischen Sinn")
    h2(doc, "Robustness und Sensitivitaet")
    p(doc, "Robustness ist die freie Stresstest-Werkbank. Sensitivitaet ist enger: Hier wird um einen optimierten Parameterwert herum variiert, um zu sehen, ob eine Strategie stabil bleibt. Der Variationskoeffizient (CV) zeigt, wie stark Ergebnisse relativ zum Mittelwert schwanken. Hohe CV-Werte deuten auf Fragilitaet hin.")
    h2(doc, "KI als Analyst, nicht als Orakel")
    p(doc, "Die KI-Auswertung liest Sensitivitaetsdaten und Performance-Kontext. Sie kann beschreiben, ob Kurven glatt, sprunghaft, peak-lastig oder plateauartig wirken. Das ist wertvoll, weil Menschen in grossen Tabellen schnell Muster uebersehen. Trotzdem ersetzt KI keine Datenvalidierung. Step 7 bleibt entscheidend.")
    p(doc, "Technisch wird hier ein LLM eingesetzt, also ein Sprachmodell, das die Ergebnisdaten als strukturierten Analysekontext bekommt. Die KI schaut sich vor allem das Verhalten der Parameter-Kennlinien an: Gibt es breite Plateaus? Fallen einzelne Parameter wie eine Klippe ab? Sind Backtest- und Forward-Sensitivitaet aehnlich? Daraus entsteht eine verbale Einschaetzung und eine Stabilitaetsbewertung mit Punkten.")
    p(doc, "Das ist hilfreich, weil Robustheit nicht immer in einer einzigen Zahl sichtbar ist. Zwei Strategien koennen denselben Profit haben, aber eine hat glatte Kennlinien und die andere nur einen schmalen Peak. Die KI kann diesen Unterschied beschreiben und damit die menschliche Auswahl beschleunigen. Sie bleibt aber Analyst, nicht Richter: Wenn Daten schlecht sind oder Step 7 scheitert, darf ein guter KI-Text den Kandidaten nicht retten.")
    image(doc, ROOT / "images" / "backtester_ki_evaluation_table.png", "KI-Bewertungstabelle: Stabilitaet, CV worst und Fazit werden verdichtet.", width=5.7)
    p(doc, "Die KI-Bewertungstabelle verdichtet viele Sensitivitaetsinformationen in eine lesbare Entscheidungshilfe. Typische Spalten zeigen Stabilitaet, schlechteste CV-Werte, Risikoindikatoren und ein kurzes Fazit. Der Wert dieser Tabelle liegt nicht darin, dass ein Sprachmodell die Zukunft kennt. Der Wert liegt darin, dass viele Kennlinien und Parameterabweichungen strukturiert beschrieben werden.")
    p(doc, "Wenn die KI zum Beispiel eine Strategie als peak-lastig markiert, sollte man in die Sensitivitaetsdetails springen und pruefen, ob Nachbarwerte wirklich stark abfallen. Wenn sie breite Plateaus erkennt, ist das ein Hinweis, aber noch kein Beweis. Die Tabelle ist also ein schneller Analystenbericht, der zur naechsten Pruefung fuehrt.")
    h2(doc, "Controlling nach dem Export")
    p(doc, "Controlling verhindert, dass eine Strategie nach dem Export aus dem Blick verschwindet. Hier werden Kennzahlen, Reviews, Nachtests und Detailanalysen zusammengefuehrt. Besonders wichtig ist die Strategie-Detailanalyse, weil sie zeigt, wie Score, Konsistenz, Backtest, Forward, Equity-Kurve und Parameter zusammenpassen.")
    p(doc, "In diesem Modul werden finale Kandidaten noch einmal intensiver betrachtet. Dazu gehoeren Nachtests auf laengeren Zeitraeumen, Real-Tick-Tests, erneute Plausibilitaetspruefungen und manuelle Reviews. Nicht alles laesst sich automatisieren: Ein Mensch muss immer noch beurteilen, ob eine Strategie fachlich sinnvoll, handelbar und zum Risiko passt. Aber die vorherigen automatisierten Schritte reduzieren die Menge drastisch. Statt tausende Paesse manuell zu lesen, schaut man im Controlling auf die wenigen Kandidaten, die bereits mehrere Gates ueberstanden haben.")
    p(doc, "Praktisch ist Controlling der Ort fuer die Frage: Wuerde ich diese Strategie wirklich weiter beobachten oder handeln? Hier sollte man besonders kritisch sein. Ein Export ist noch kein Ritterschlag, sondern nur der Eintritt in die naechste Qualitaetsstufe.")
    image(doc, ROOT / "images" / "backtester_strategy_detail_analysis.png", "Strategie-Detailanalyse: Kennzahlen, Score-Erklaerung, Equity-Kurve und Parameter im Zusammenhang.", width=6.35)
    p(doc, "Die Strategie-Detailanalyse verbindet die wichtigsten Entscheidungsebenen an einem Ort. Man sieht Kennzahlen, Score-Erklaerung, Equity-Kurve und Parameter nicht getrennt, sondern im Zusammenhang. Genau das verhindert Fehlurteile: Ein hoher Score ist nur dann wertvoll, wenn die Equity-Kurve nachvollziehbar ist, der Drawdown zum Kontorisiko passt, genug Trades vorhanden sind und die Parameter nicht nach einem zufaelligen Peak aussehen.")
    p(doc, "In der Praxis ist diese Ansicht die letzte Plausibilitaetsbruecke vor weiteren Nachtests. Sie beantwortet: Warum wurde dieser Kandidat ausgewaehlt? Welche Schwachstellen sind noch sichtbar? Welche Zusatztests sind sinnvoll - laengerer Zeitraum, Real-Tick-Test, anderer Brokerdatenstand oder erneute OOS-Pruefung?")
    source_note(doc, "Walk-Forward-Optimierung reduziert Overfitting, indem Parameter wiederholt in nachfolgenden Daten validiert werden; sie bleibt aber von Fensterwahl, Regimewechseln und Rechenaufwand abhaengig [S5].")


def add_supporting_areas_section(doc: Document):
    h1(doc, "8. Weitere Arbeitsbereiche: Daten, Settings, Log, Database und Manual")
    p(doc, "Die Nebenbereiche sind keine Dekoration. Sie entscheiden, ob Backtests reproduzierbar, nachvollziehbar und fehlerdiagnostizierbar sind. Viele scheinbare Strategieprobleme sind in Wahrheit Daten-, Pfad-, Zeitzonen- oder Reportprobleme.")
    p(doc, "Ein zentrales Prinzip des Projekts ist Persistenz: Die wichtigen Schritte werden in Datenbanken festgehalten. Backtests, Optimierungen, Workflow-Zustand, KI-Berichte, Parameter-Snapshots und Historien bleiben damit spaeter abrufbar. Das ist ein grosser Vorteil, wenn man in einem Jahr dieselbe Strategie erneut optimieren will. Man sieht dann, welche Parameter frueher funktionierten, welche Filter gesetzt waren, welche Datenfenster benutzt wurden und warum ein Kandidat akzeptiert oder verworfen wurde.")
    p(doc, "Datenbankgestuetzte Arbeit schuetzt auch vor Bauchgefuehl. Ohne Historie erinnert man sich oft nur an die besten Ergebnisse. Mit Historie sieht man auch die verworfenen Runs, die Fehlstarts und die Marktphasen, in denen eine Idee nicht funktioniert hat. Genau diese negativen Informationen sind fuer robuste Strategieentwicklung wertvoll.")
    h2(doc, "Settings: die technische Grundlage")
    rows = [
        ["MT5 Terminal Path", "Pfad zu terminal64.exe.", "Ohne korrekten Pfad kann kein MT5-Test gestartet werden."],
        ["MT4 Terminal Path", "Pfad zu terminal.exe.", "Wichtig fuer alte MT4-EAs oder Vergleichslaeufe."],
        ["Portable Mode", "Startet MetaTrader mit /portable.", "Sorgt dafuer, dass Daten und Konfiguration im Terminalordner bleiben."],
        ["Reports Output", "Ablageort fuer XML/HTML/Reports.", "Muss beschreibbar und gut auffindbar sein."],
        ["Data Directory", "Arbeitsordner fuer Marktdaten.", "Nicht mit Reportordnern vermischen."],
        ["Default Deposit / Currency / Leverage", "Standardkonto fuer neue Laeufe.", "Als Laborstandard definieren, damit Ergebnisse vergleichbar bleiben."],
        ["Default Tick Model", "Vorgabe fuer neue Tests.", "Schnelle Vorpruefung und finale Pruefung bewusst unterscheiden."],
        ["Broker Timezone", "UTC-Offset des Broker-Servers.", "Wichtig bei Datenimport, Tageswechsel und Session-Effekten."],
    ]
    table(doc, ["Setting", "Funktion", "Warum wichtig"], rows, [1.65, 2.35, 2.5], "Application Settings")

    h2(doc, "Dukascopy Data: unabhängige Datenbasis")
    p(doc, "Dukascopy liefert historische Tickdaten als BI5-Dateien. Das Projekt kann diese Daten herunterladen, scannen, in CSV/M1-Strukturen konvertieren und fuer MT5 Custom Symbols vorbereiten. Der Sinn ist Datenkontrolle: Wer nur Broker-Historie nutzt, uebernimmt auch deren Luecken, Session-Regeln und Spreadannahmen.")
    p(doc, "Dukascopy-Daten sind keine Pflicht, aber eine starke Ergaenzung. Dukascopy Bank SA ist eine Schweizer Online-Bank und Forex-/CFD-Anbieterin mit Sitz in Genf; die offiziellen Seiten beschreiben Historical Data Export und historische Tick-Abfragen im JForex/API-Kontext [S9, S10, S11]. In der Praxis werden Dukascopy-Daten haeufig genutzt, weil sie detaillierte Historien und Tickdaten fuer viele Instrumente bereitstellen. Fuer Backtesting ist das wertvoll, weil man nicht nur von der Datenqualitaet eines einzelnen MetaTrader-Brokers abhaengig ist.")
    p(doc, "Der wichtigste Nutzen ist der Datenvergleich. Eine Strategie, die nur auf der Historie eines Brokers gut aussieht, kann an speziellen Kursluecken, Spreads, Zeitzonen oder Datenfehlern haengen. Wenn sie dagegen auch auf einer zweiten Datenquelle wie Dukascopy plausibel bleibt, steigt das Vertrauen. Das beweist keine Live-Performance, aber es reduziert die Gefahr, dass man eine Strategie nur auf eine einzige Datenreihe angepasst hat.")
    callout(
        doc,
        "Einfaches Beispiel: Brokerdaten gegen Dukascopy",
        "Ein EA gewinnt auf Broker A im Zeitraum 2022 bis 2024 solide. Danach wird derselbe Zeitraum mit importierten Dukascopy-Daten als MT5 Custom Symbol getestet. Bleiben Trades, Drawdown und Equity-Verlauf aehnlich, ist das ein gutes Robustheitssignal. Bricht die Strategie komplett ein, muss man klaeren, ob Spread, Zeitzone, Tickqualitaet oder ein echter Strategiefehler die Ursache ist.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    rows = [
        ["Select All / Select None", "Markiert Symbolgruppen fuer Downloads.", "Nur Daten laden, die wirklich gebraucht werden."],
        ["Download Data", "Laedt BI5-Daten fuer gewaehlte Symbole und Zeitraeume.", "Auf Speicherplatz und Laufzeit achten."],
        ["Scan Downloaded Data", "Prueft, welche Daten lokal vorhanden sind.", "Vor Konvertierung nutzen, um Luecken zu erkennen."],
        ["Convert to CSV", "Wandelt BI5/Tickdaten in verarbeitbare CSV-Strukturen.", "Konvertierung nach Datenpruefung starten."],
        ["Import to MT5", "Bereitet Import in MT5 Custom Symbols vor.", "Danach Symbol in MetaTrader verifizieren."],
        ["Export CSV", "Exportiert Daten fuer externe Analyse.", "Gut fuer Plausibilitaetschecks in Python/Excel."],
    ]
    table(doc, ["Aktion", "Funktion", "Best Practice"], rows, [1.55, 2.45, 2.5], "Dukascopy Data")

    h2(doc, "Database und History: Nachvollziehbarkeit statt Zettelwirtschaft")
    p(doc, "Die SQLite-Persistenz speichert Backtests, Optimierungen, Workflow-Zustände, KI-Berichte und Parameterstände. Das ist wichtig, weil robuste Strategieentwicklung nicht aus einem einzigen Run besteht. Man muss spaeter nachvollziehen koennen, welche Annahmen zu welchem Kandidaten gefuehrt haben.")
    p(doc, "Der Vorteil zeigt sich besonders spaeter. Wenn man nach Monaten oder nach einem Jahr erneut optimiert, muss man nicht bei null beginnen. Man kann alte Ergebnisse vergleichen, alte SET-Dateien laden, pruefen, ob die neue Marktphase andere Parameter bevorzugt, und erkennen, ob sich eine Strategie verschlechtert oder nur temporaer anders verhaelt. Die Datenbank wird damit zum Forschungstagebuch des Projekts.")
    rows = [
        ["Backtest History", "Liste einzelner Testlaeufe.", "Bei Regressionen nachsehen, wann ein EA noch funktionierte."],
        ["Optimization Results", "Optimierungspasses und Combined-Daten.", "Filterentscheidungen nicht nur im Kopf behalten."],
        ["Saved Configs", "Parameter-Snapshots pro EA.", "Saubere Versionierung von SET-Ideen."],
        ["KI Reports", "Gespeicherte LLM-Auswertungen.", "KI-Fazit mit echten Kennzahlen zusammen lesen."],
        ["Workflow State", "Status der Pipeline.", "Erkennt, ob Step 7 wirklich abgeschlossen wurde."],
    ]
    table(doc, ["Datenbereich", "Speichert", "Nutzen"], rows, [1.55, 2.35, 2.6], "Database / History")

    h2(doc, "Log: die Wahrheit über den Lauf")
    p(doc, "Der Log-Bereich zeigt, was das Programm im Hintergrund tut: Pfade, Startbefehle, Prozessstatus, Parsermeldungen, Fehler und Warnungen. Wenn ein Test nicht startet, kein Report entsteht oder MT5 offen bleibt, ist der Log der erste Ort. Er trennt Bedienfehler, Datenprobleme, Terminalprobleme und echte Strategieprobleme.")
    callout(
        doc,
        "Professionelle Fehlersuche",
        "Nicht zuerst Parameter aendern, wenn ein Lauf fehlschlaegt. Zuerst Log lesen: Terminalpfad, EA-Pfad, Symbolverfuegbarkeit, Reportordner, Rechte, laufende MT5-Instanzen und Parsermeldungen pruefen.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    h2(doc, "Manual: Schnellhilfe innerhalb der Anwendung")
    p(doc, "Der Manual-Tab ersetzt kein vollstaendiges Buch, ist aber als Erinnerung im Arbeitsfluss nuetzlich. Dieses Dokument erklaert Hintergruende und Entscheidungen; die In-App-Hilfe ist fuer schnelles Nachschlagen gedacht, wenn man gerade im Programm arbeitet.")


def add_reference_section(doc: Document):
    h1(doc, "9. Filter, Parameter und Best Practices")
    h2(doc, "Best Practices fuer robuste Optimierung")
    bullet(doc, "Vor der Optimierung entscheiden, welches spaetere Zeitfenster wirklich out-of-sample bleiben soll.")
    bullet(doc, "Nur Parameter optimieren, die eine fachliche Bedeutung haben.")
    bullet(doc, "Forward-Test aktivieren und Forward-Ergebnisse nicht als finale Wahrheit missverstehen.")
    bullet(doc, "Mindesttrades ernst nehmen: Wenige Trades sind schwache Statistik.")
    bullet(doc, "Drawdown, Recovery und Konsistenz hoeher gewichten als reine Maximalprofite.")
    bullet(doc, "Sensitivitaetskurven visuell lesen: Plateaus sind besser als Peaks.")
    bullet(doc, "Multi-Backtester als Screening nutzen, nicht als Lotterie nach dem besten Einzelrun.")
    bullet(doc, "Step-7-Validierung nicht ueberspringen, wenn es um reale Kandidaten geht.")

    h2(doc, "Was Backtester und Multi-Backtester nicht tun")
    p(doc, "Diese beiden Bereiche bauen keine eigene Handelslogik und keine neue Simulationsengine. Sie verwenden die Parameter, die MetaTrader ohnehin kennt. Der Unterschied liegt in der Automatisierung: Statt jeden Lauf manuell einzurichten, zu starten und den Report händisch zu lesen, erzeugt das Programm Konfigurationen, startet MetaTrader kontrolliert und sammelt Kennzahlen automatisch.")
    h2(doc, "Was der Optimizer zusaetzlich leistet")
    p(doc, "Der Optimizer nutzt zwar MetaTrader-Optimierung als Grundlage, aber die eigentliche Qualitaet entsteht in der zusaetzlichen Auswertung. Combined Analysis, Filter, Score, Sensitivitaet und KI-Bericht sorgen dafuer, dass nicht nur der hoechste Profit gewinnt. Besonders das Plateau-Denken ist der Unterschied: Gesucht wird eine stabile Region, nicht die schoenste Spitze.")
    callout(
        doc,
        "Professioneller Entscheidungsstandard",
        "Eine Strategie ist erst dann interessant, wenn sie technische Plausibilitaet, ausreichende Stichprobe, kontrolliertes Risiko, Forward-Konsistenz, Sensitivitaetsstabilitaet und eine spaetere OOS-Pruefung zusammenbringt.",
        fill=LIGHT_GRAY,
        accent=BLUE,
    )


def add_sources(doc: Document):
    doc.add_page_break()
    h1(doc, "10. Quellenverzeichnis")
    p(doc, "Die folgenden Quellen wurden fuer die fachlichen Grundlagen zu MetaTrader, Backtesting, Optimierung, Overfitting, Parameter-Plateaus, Forex-Maerkten und Timeframes verwendet. Der Text im Handbuch ist als eigene Zusammenfassung formuliert; externe Inhalte werden nicht laenglich kopiert.")
    rows = [[sid, org, title, url] for sid, org, title, url in SOURCES]
    table(doc, ["ID", "Quelle", "Titel", "URL"], rows, [0.45, 1.35, 2.25, 2.45])


def build_docx():
    os.chdir(ROOT)
    graphics = make_manual_graphics()
    doc = Document()
    configure_document(doc)
    add_cover(doc, graphics)
    add_toc(doc)
    add_frontmatter(doc)
    add_backtesting_foundation(doc, graphics)
    doc.add_page_break()
    add_program_overview(doc)
    doc.add_page_break()
    add_backtester_section(doc)
    doc.add_page_break()
    add_multibacktester_section(doc, graphics)
    doc.add_page_break()
    add_optimizer_section(doc, graphics)
    doc.add_page_break()
    add_workflow_section(doc)
    doc.add_page_break()
    add_supporting_areas_section(doc)
    doc.add_page_break()
    add_reference_section(doc)
    add_sources(doc)
    try:
        doc.save(DOCX_PATH)
        print(f"Wrote {DOCX_PATH}")
    except PermissionError:
        doc.save(FALLBACK_DOCX_PATH)
        print(f"Wrote {FALLBACK_DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
